from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
import base64
import json
import re
import xml.etree.ElementTree as ET
from xml.dom import minidom
import os
import tempfile
import shutil
import subprocess

# Check for optional dependencies
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

try:
    import markdown
    import html2text
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def index(request):
    """Text converters main page"""
    return render(request, 'text_converter/index.html', {
        'yaml_available': YAML_AVAILABLE,
        'markdown_available': MARKDOWN_AVAILABLE,
    })


def uppercase_lowercase(request):
    """Convert text between uppercase and lowercase"""
    if request.method != 'POST':
        return render(request, 'text_converter/uppercase_lowercase.html')
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'to_upper')
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/uppercase_lowercase.html')
    
    if conversion_type == 'to_upper':
        result = text.upper()
    elif conversion_type == 'to_lower':
        result = text.lower()
    elif conversion_type == 'to_title':
        result = text.title()
    elif conversion_type == 'to_sentence':
        result = text.capitalize()
    else:
        result = text
    
    return render(request, 'text_converter/uppercase_lowercase.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
    })


def camelcase_snakecase(request):
    """Convert between CamelCase and snake_case"""
    if request.method != 'POST':
        return render(request, 'text_converter/camelcase_snakecase.html')
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'to_snake')
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/camelcase_snakecase.html')
    
    if conversion_type == 'to_snake':
        # Convert CamelCase to snake_case
        # Insert underscore before uppercase letters (except first)
        result = re.sub(r'(?<!^)(?=[A-Z])', '_', text).lower()
    elif conversion_type == 'to_camel':
        # Convert snake_case to CamelCase
        components = text.split('_')
        result = components[0].lower() + ''.join(word.capitalize() for word in components[1:])
    elif conversion_type == 'to_pascal':
        # Convert snake_case to PascalCase
        components = text.split('_')
        result = ''.join(word.capitalize() for word in components)
    else:
        result = text
    
    return render(request, 'text_converter/camelcase_snakecase.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
    })


def remove_special_characters(request):
    """Remove special characters from text"""
    if request.method != 'POST':
        return render(request, 'text_converter/remove_special.html')
    
    text = request.POST.get('text', '').strip()
    keep_spaces = request.POST.get('keep_spaces', 'on') == 'on'
    keep_numbers = request.POST.get('keep_numbers', 'on') == 'on'
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/remove_special.html')
    
    # Build pattern: keep letters, optionally spaces and numbers
    if keep_spaces and keep_numbers:
        pattern = r'[^a-zA-Z0-9\s]'
    elif keep_spaces:
        pattern = r'[^a-zA-Z\s]'
    elif keep_numbers:
        pattern = r'[^a-zA-Z0-9]'
    else:
        pattern = r'[^a-zA-Z]'
    
    result = re.sub(pattern, '', text)
    
    return render(request, 'text_converter/remove_special.html', {
        'input_text': text,
        'result': result,
        'keep_spaces': keep_spaces,
        'keep_numbers': keep_numbers,
    })


def remove_duplicate_lines(request):
    """Remove duplicate lines from text"""
    if request.method != 'POST':
        return render(request, 'text_converter/remove_duplicates.html')
    
    text = request.POST.get('text', '').strip()
    preserve_order = request.POST.get('preserve_order', 'on') == 'on'
    case_sensitive = request.POST.get('case_sensitive', 'on') == 'on'
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/remove_duplicates.html')
    
    lines = text.split('\n')
    
    if preserve_order:
        seen = set()
        result_lines = []
        for line in lines:
            key = line if case_sensitive else line.lower()
            if key not in seen:
                seen.add(key)
                result_lines.append(line)
        result = '\n'.join(result_lines)
    else:
        if case_sensitive:
            unique_lines = list(dict.fromkeys(lines))
        else:
            seen = set()
            unique_lines = []
            for line in lines:
                key = line.lower()
                if key not in seen:
                    seen.add(key)
                    unique_lines.append(line)
        result = '\n'.join(unique_lines)
    
    removed_count = len(lines) - len(result.split('\n'))
    
    return render(request, 'text_converter/remove_duplicates.html', {
        'input_text': text,
        'result': result,
        'removed_count': removed_count,
        'preserve_order': preserve_order,
        'case_sensitive': case_sensitive,
    })


def sort_lines(request):
    """Sort lines alphabetically"""
    if request.method != 'POST':
        return render(request, 'text_converter/sort_lines.html')
    
    text = request.POST.get('text', '').strip()
    sort_order = request.POST.get('sort_order', 'asc')
    case_sensitive = request.POST.get('case_sensitive', 'on') == 'on'
    remove_empty = request.POST.get('remove_empty', 'off') == 'on'
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/sort_lines.html')
    
    lines = text.split('\n')
    
    if remove_empty:
        lines = [line for line in lines if line.strip()]
    
    if case_sensitive:
        sorted_lines = sorted(lines, reverse=(sort_order == 'desc'))
    else:
        sorted_lines = sorted(lines, key=str.lower, reverse=(sort_order == 'desc'))
    
    result = '\n'.join(sorted_lines)
    
    return render(request, 'text_converter/sort_lines.html', {
        'input_text': text,
        'result': result,
        'sort_order': sort_order,
        'case_sensitive': case_sensitive,
        'remove_empty': remove_empty,
    })


def json_xml(request):
    """Convert between JSON and XML"""
    if request.method != 'POST':
        return render(request, 'text_converter/json_xml.html')
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'json_to_xml')
    
    if not text:
        messages.error(request, 'Please enter JSON or XML text.')
        return render(request, 'text_converter/json_xml.html')
    
    try:
        if conversion_type == 'json_to_xml':
            # Parse JSON
            data = json.loads(text)
            # Convert to XML
            root = ET.Element('root')
            _dict_to_xml(data, root)
            # Format XML
            xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
            result = '\n'.join(line for line in xml_str.split('\n') if line.strip())
        else:  # xml_to_json
            # Parse XML
            root = ET.fromstring(text)
            # Convert to dict
            data = _xml_to_dict(root)
            # Convert to JSON
            result = json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as e:
        messages.error(request, f'Invalid JSON: {str(e)}')
        return render(request, 'text_converter/json_xml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
        })
    except ET.ParseError as e:
        messages.error(request, f'Invalid XML: {str(e)}')
        return render(request, 'text_converter/json_xml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
        })
    except Exception as e:
        messages.error(request, f'Conversion error: {str(e)}')
        return render(request, 'text_converter/json_xml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
        })
    
    return render(request, 'text_converter/json_xml.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
    })


def _dict_to_xml(d, root):
    """Helper function to convert dict to XML"""
    for key, value in d.items():
        child = ET.SubElement(root, str(key))
        if isinstance(value, dict):
            _dict_to_xml(value, child)
        elif isinstance(value, list):
            for item in value:
                item_elem = ET.SubElement(child, 'item')
                if isinstance(item, dict):
                    _dict_to_xml(item, item_elem)
                else:
                    item_elem.text = str(item)
        else:
            child.text = str(value)


def _xml_to_dict(element):
    """Helper function to convert XML to dict"""
    result = {}
    for child in element:
        if len(child) == 0:
            # Leaf node - check if tag already exists (multiple children with same tag)
            if child.tag in result:
                # Convert to list if not already
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child.text)
            else:
                result[child.tag] = child.text
        else:
            # Has children - recursively convert
            child_dict = _xml_to_dict(child)
            if child.tag in result:
                # Convert to list if not already
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_dict)
            else:
                result[child.tag] = child_dict
    return result if result else None


def json_yaml(request):
    """Convert between JSON and YAML"""
    if request.method != 'POST':
        return render(request, 'text_converter/json_yaml.html', {
            'yaml_available': YAML_AVAILABLE,
        })
    
    if not YAML_AVAILABLE:
        messages.error(request, 'YAML support requires PyYAML library. Install with: pip install pyyaml')
        return render(request, 'text_converter/json_yaml.html', {
            'yaml_available': YAML_AVAILABLE,
        })
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'json_to_yaml')
    
    if not text:
        messages.error(request, 'Please enter JSON or YAML text.')
        return render(request, 'text_converter/json_yaml.html', {
            'yaml_available': YAML_AVAILABLE,
        })
    
    try:
        if conversion_type == 'json_to_yaml':
            # Parse JSON
            data = json.loads(text)
            # Convert to YAML
            result = yaml.dump(data, default_flow_style=False, allow_unicode=True, sort_keys=False)
        else:  # yaml_to_json
            # Parse YAML
            data = yaml.safe_load(text)
            # Convert to JSON
            result = json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError as e:
        messages.error(request, f'Invalid JSON: {str(e)}')
        return render(request, 'text_converter/json_yaml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
            'yaml_available': YAML_AVAILABLE,
        })
    except yaml.YAMLError as e:
        messages.error(request, f'Invalid YAML: {str(e)}')
        return render(request, 'text_converter/json_yaml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
            'yaml_available': YAML_AVAILABLE,
        })
    except Exception as e:
        messages.error(request, f'Conversion error: {str(e)}')
        return render(request, 'text_converter/json_yaml.html', {
            'input_text': text,
            'conversion_type': conversion_type,
            'yaml_available': YAML_AVAILABLE,
        })
    
    return render(request, 'text_converter/json_yaml.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
        'yaml_available': YAML_AVAILABLE,
    })


def html_markdown(request):
    """Convert between HTML and Markdown"""
    if request.method != 'POST':
        return render(request, 'text_converter/html_markdown.html', {
            'markdown_available': MARKDOWN_AVAILABLE,
        })
    
    if not MARKDOWN_AVAILABLE:
        messages.error(request, 'Markdown support requires markdown and html2text libraries. Install with: pip install markdown html2text')
        return render(request, 'text_converter/html_markdown.html', {
            'markdown_available': MARKDOWN_AVAILABLE,
        })
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'html_to_markdown')
    
    if not text:
        messages.error(request, 'Please enter HTML or Markdown text.')
        return render(request, 'text_converter/html_markdown.html', {
            'markdown_available': MARKDOWN_AVAILABLE,
        })
    
    try:
        if conversion_type == 'html_to_markdown':
            # Convert HTML to Markdown
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            h.body_width = 0  # Don't wrap lines
            result = h.handle(text).strip()
        else:  # markdown_to_html
            # Convert Markdown to HTML
            result = markdown.markdown(text, extensions=['extra', 'codehilite'])
    except Exception as e:
        messages.error(request, f'Conversion error: {str(e)}')
        return render(request, 'text_converter/html_markdown.html', {
            'input_text': text,
            'conversion_type': conversion_type,
            'markdown_available': MARKDOWN_AVAILABLE,
        })
    
    return render(request, 'text_converter/html_markdown.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
        'markdown_available': MARKDOWN_AVAILABLE,
    })


def text_base64(request):
    """Convert between text and Base64"""
    if request.method != 'POST':
        return render(request, 'text_converter/text_base64.html')
    
    text = request.POST.get('text', '').strip()
    conversion_type = request.POST.get('conversion_type', 'text_to_base64')
    
    if not text:
        messages.error(request, 'Please enter some text.')
        return render(request, 'text_converter/text_base64.html')
    
    try:
        if conversion_type == 'text_to_base64':
            # Convert text to Base64
            text_bytes = text.encode('utf-8')
            result = base64.b64encode(text_bytes).decode('utf-8')
        else:  # base64_to_text
            # Convert Base64 to text
            text_bytes = base64.b64decode(text)
            result = text_bytes.decode('utf-8')
    except Exception as e:
        messages.error(request, f'Conversion error: {str(e)}. Please ensure the input is valid.')
        return render(request, 'text_converter/text_base64.html', {
            'input_text': text,
            'conversion_type': conversion_type,
        })
    
    return render(request, 'text_converter/text_base64.html', {
        'input_text': text,
        'result': result,
        'conversion_type': conversion_type,
    })


def word_counter(request):
    """Count words, characters, lines, etc."""
    if request.method != 'POST':
        return render(request, 'text_converter/word_counter.html')
    
    text = request.POST.get('text', '').strip()
    
    if not text:
        # Still show the page with zero counts
        return render(request, 'text_converter/word_counter.html', {
            'input_text': '',
            'char_count': 0,
            'char_count_no_spaces': 0,
            'word_count': 0,
            'line_count': 0,
            'paragraph_count': 0,
            'sentence_count': 0,
        })
    
    # Count characters
    char_count = len(text)
    char_count_no_spaces = len(text.replace(' ', '').replace('\n', '').replace('\t', ''))
    
    # Count words
    words = text.split()
    word_count = len(words)
    
    # Count lines
    lines = text.split('\n')
    line_count = len(lines)
    
    # Count paragraphs (non-empty lines separated by blank lines)
    paragraphs = [p for p in text.split('\n\n') if p.strip()]
    paragraph_count = len(paragraphs)
    
    # Count sentences (rough estimate)
    sentence_endings = re.findall(r'[.!?]+', text)
    sentence_count = len(sentence_endings) if sentence_endings else 1
    
    return render(request, 'text_converter/word_counter.html', {
        'input_text': text,
        'char_count': char_count,
        'char_count_no_spaces': char_count_no_spaces,
        'word_count': word_count,
        'line_count': line_count,
        'paragraph_count': paragraph_count,
        'sentence_count': sentence_count,
    })


def _check_ffprobe():
    """Check if ffprobe is available and return its path"""
    ffprobe_path = shutil.which('ffprobe') or '/usr/local/bin/ffprobe'
    
    if not os.path.exists(ffprobe_path):
        return None, 'ffprobe not found. Please install FFmpeg on your system.'
    
    ffprobe_dir = os.path.dirname(ffprobe_path)
    original_path = os.environ.get('PATH', '')
    if ffprobe_dir not in original_path:
        os.environ['PATH'] = f"{ffprobe_dir}:{original_path}"
    
    return ffprobe_path, None


def _get_audio_duration(file_path, ffprobe_path):
    """
    Get audio file duration in seconds using ffprobe.
    Returns duration as float, or None if error.
    """
    try:
        cmd = [
            ffprobe_path,
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'default=noprint_wrappers=1:nokey=1',
            file_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode != 0:
            return None
        
        duration_str = result.stdout.strip()
        if duration_str:
            return float(duration_str)
        return None
    except (ValueError, subprocess.TimeoutExpired, Exception):
        return None


def _transcribe_audio(file_path, api_key):
    """
    Transcribe audio file using OpenAI Whisper API.
    Returns tuple: (transcription_text, detected_language) or (None, None) on error.
    """
    try:
        # Create client with extended timeout for large audio files (up to 10 minutes)
        client = OpenAI(
            api_key=api_key,
            timeout=600.0  # 10 minutes timeout for large files
        )
        
        with open(file_path, 'rb') as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language=None  # Auto-detect language (do NOT hardcode)
            )
        
        transcription_text = transcript.text
        # Whisper API response may include language if available
        # Check for language attribute in the response
        detected_language = getattr(transcript, 'language', None)
        
        return transcription_text, detected_language
    except Exception as e:
        # Log error for debugging but don't expose internal details to user
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f'OpenAI transcription error: {str(e)}')
        return None, None


def audio_transcription(request):
    """Transcribe audio files using OpenAI Whisper API"""
    if request.method != 'POST':
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    # Check if OpenAI is available
    if not OPENAI_AVAILABLE:
        messages.error(request, 'OpenAI library is not installed. Please install with: pip install openai')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    # Check for API key
    from django.conf import settings
    api_key = settings.OPENAI_API_KEY
    if not api_key:
        messages.error(request, 'OpenAI API key is not configured. Please set OPENAI_API_KEY in your .env file.')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    # Check for uploaded file
    if 'audio_file' not in request.FILES:
        messages.error(request, 'Please upload an audio file.')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    uploaded_file = request.FILES['audio_file']
    
    # Validate file size (max 700MB, consistent with other tools)
    max_size = 700 * 1024 * 1024  # 700MB
    if uploaded_file.size > max_size:
        messages.error(request, f'File size exceeds 700MB limit. Your file is {uploaded_file.size / (1024 * 1024):.1f}MB.')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    # Validate file type
    allowed_extensions = ['.mp3', '.wav', '.m4a', '.ogg']
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_ext not in allowed_extensions:
        messages.error(request, 'Invalid file type. Please upload MP3, WAV, M4A, or OGG files.')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    
    # Check ffprobe for duration validation
    ffprobe_path, error = _check_ffprobe()
    if error:
        # If ffprobe is not available, we'll skip server-side duration check
        # but still proceed (client-side validation should catch it)
        ffprobe_path = None
    
    temp_dir = None
    try:
        # Create temporary directory
        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, f'input{file_ext}')
        
        # Save uploaded file
        with open(input_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Validate audio duration using ffprobe (server-side check)
        if ffprobe_path:
            duration = _get_audio_duration(input_path, ffprobe_path)
            if duration is not None:
                max_duration = 30 * 60  # 30 minutes in seconds
                if duration > max_duration:
                    messages.error(request, f'Audio duration ({duration / 60:.1f} minutes) exceeds the 30-minute limit. Please use the audio splitter tool to shorten your file.')
                    return render(request, 'text_converter/audio_transcription.html', {
                        'openai_available': OPENAI_AVAILABLE,
                        'docx_available': DOCX_AVAILABLE,
                        'duration_exceeded': True,
                    })
        
        # Transcribe audio using OpenAI Whisper API
        transcription_text, detected_language = _transcribe_audio(input_path, api_key)
        
        if transcription_text is None:
            messages.error(request, 'Failed to transcribe audio. Please check your API key and try again.')
            return render(request, 'text_converter/audio_transcription.html', {
                'openai_available': OPENAI_AVAILABLE,
                'docx_available': DOCX_AVAILABLE,
            })
        
        # Return results
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'transcription': transcription_text,
            'detected_language': detected_language,
            'docx_available': DOCX_AVAILABLE,
        })
        
    except Exception as e:
        messages.error(request, f'Error processing file: {str(e)}')
        return render(request, 'text_converter/audio_transcription.html', {
            'openai_available': OPENAI_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
        })
    finally:
        # Clean up temporary files
        if temp_dir:
            try:
                shutil.rmtree(temp_dir)
            except (OSError, PermissionError):
                pass


def download_transcription_docx(request):
    """Download transcription as .docx file"""
    if request.method != 'POST':
        return HttpResponse('Method not allowed', status=405)
    
    if not DOCX_AVAILABLE:
        return HttpResponse('python-docx library is not installed', status=503)
    
    transcription_text = request.POST.get('transcription', '').strip()
    
    if not transcription_text:
        return HttpResponse('No transcription text provided', status=400)
    
    try:
        import io
        from datetime import datetime
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('Audio Transcription', 0)
        
        # Add transcription text
        # Split by paragraphs (double newlines) for better formatting
        # If no double newlines, split by single newlines
        if '\n\n' in transcription_text:
            paragraphs = transcription_text.split('\n\n')
        else:
            # Split by single newlines if no double newlines found
            paragraphs = transcription_text.split('\n')
        
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to BytesIO buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')
        filename = f'transcription-{timestamp}.docx'
        
        # Create HTTP response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f'Error creating DOCX: {str(e)}')
        return HttpResponse(f'Error creating document: {str(e)}', status=500)

