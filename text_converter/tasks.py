"""
Celery tasks for text converter operations
"""
import os
import tempfile
import shutil
import subprocess
import logging
from datetime import datetime
from celery import shared_task
from django.conf import settings
from django.utils import timezone
from openai import OpenAI
from openai import RateLimitError as OpenAIRateLimitError
from openai import APIError as OpenAIAPIError
from .models import TranscriptionJob

logger = logging.getLogger(__name__)


def _check_ffprobe():
    """Check if ffprobe is available"""
    try:
        result = subprocess.run(['which', 'ffprobe'], capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            return result.stdout.strip(), None
        return None, "ffprobe not found in PATH"
    except Exception as e:
        return None, str(e)


def _get_audio_duration(file_path, ffprobe_path):
    """Get audio duration in seconds using ffprobe"""
    try:
        cmd = [
            ffprobe_path,
            '-v', 'error',
            '-show_entries', 'format=duration',
            '-of', 'default=noprint_wrappers=1:nokey=1',
            file_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return float(result.stdout.strip())
    except (ValueError, subprocess.TimeoutExpired, Exception):
        return None
    return None


def estimate_transcription_cost(file_size_mb, audio_duration_minutes=None):
    """
    Estimate transcription cost.
    
    Args:
        file_size_mb: File size in MB
        audio_duration_minutes: Audio duration in minutes (if known)
    
    Returns:
        Estimated cost in USD
    """
    if audio_duration_minutes:
        # Direct calculation from duration
        return audio_duration_minutes * 0.006
    else:
        # Estimate: 1MB ≈ 1 minute for typical audio (MP3 128kbps)
        estimated_minutes = file_size_mb * 0.8  # Conservative estimate
        return estimated_minutes * 0.006


class WhisperAPICircuitBreaker:
    """Circuit breaker for OpenAI Whisper API to prevent cascading failures"""
    FAILURE_THRESHOLD = 0.5  # 50% failure rate
    TIME_WINDOW = 300  # 5 minutes
    CIRCUIT_OPEN_DURATION = 600  # 10 minutes
    
    @classmethod
    def is_open(cls):
        """Check if circuit breaker is open"""
        from django.core.cache import cache
        return cache.get('whisper_circuit_open', False)
    
    @classmethod
    def record_success(cls):
        """Record successful API call"""
        from django.core.cache import cache
        cache.delete('whisper_circuit_open')
        # Reset failure count
        cache.set('whisper_failures', 0, cls.TIME_WINDOW)
    
    @classmethod
    def record_failure(cls):
        """Record failed API call and check threshold"""
        from django.core.cache import cache
        failures = cache.get('whisper_failures', 0) + 1
        total = cache.get('whisper_total', 0) + 1
        
        cache.set('whisper_failures', failures, cls.TIME_WINDOW)
        cache.set('whisper_total', total, cls.TIME_WINDOW)
        
        failure_rate = failures / total if total > 0 else 0
        
        if failure_rate >= cls.FAILURE_THRESHOLD and total >= 10:
            # Open circuit breaker
            cache.set('whisper_circuit_open', True, cls.CIRCUIT_OPEN_DURATION)
            return True
        return False


@shared_task(
    bind=True,
    soft_time_limit=900,  # 15 minutes
    time_limit=960,  # 16 minutes
    autoretry_for=(OpenAIRateLimitError, OpenAIAPIError),
    retry_backoff=True,
    retry_backoff_max=600,
    max_retries=3,
    queue='transcription'
)
def transcribe_audio_task(self, job_id, file_path, api_key):
    """
    Transcribe audio file using OpenAI Whisper API.
    
    Args:
        job_id: UUID of the TranscriptionJob
        file_path: Path to the audio file
        api_key: OpenAI API key
    
    Returns:
        dict with transcription_text and detected_language
    """
    job = TranscriptionJob.objects.get(job_id=job_id)
    
    # Check circuit breaker
    if WhisperAPICircuitBreaker.is_open():
        error_msg = "OpenAI API circuit breaker is open. Please try again later."
        job.status = 'failed'
        job.error_message = error_msg
        job.completed_at = timezone.now()
        job.save()
        logger.error(f"Circuit breaker open for job {job_id}")
        raise Exception(error_msg)
    
    try:
        # Update job status
        job.status = 'processing'
        job.save()
        
        # Create client with extended timeout
        client = OpenAI(
            api_key=api_key,
            timeout=600.0  # 10 minutes timeout for large files
        )
        
        # Transcribe audio
        with open(file_path, 'rb') as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language=None  # Auto-detect language
            )
        
        transcription_text = transcript.text
        detected_language = getattr(transcript, 'language', None)
        
        # Calculate actual cost
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        duration_minutes = None
        if job.duration_seconds:
            duration_minutes = job.duration_seconds / 60
        
        actual_cost = estimate_transcription_cost(file_size_mb, duration_minutes)
        
        # Update job with results
        job.status = 'completed'
        job.transcription_text = transcription_text
        job.detected_language = detected_language
        job.actual_cost = actual_cost
        job.completed_at = timezone.now()
        job.save()
        
        # Record success
        WhisperAPICircuitBreaker.record_success()
        
        logger.info(f"Transcription completed for job {job_id}")
        
        return {
            'transcription_text': transcription_text,
            'detected_language': detected_language
        }
        
    except OpenAIRateLimitError as e:
        # Rate limit error - retry with exponential backoff
        error_msg = f"OpenAI API rate limit exceeded: {str(e)}"
        logger.warning(f"Rate limit for job {job_id}: {error_msg}")
        WhisperAPICircuitBreaker.record_failure()
        raise self.retry(exc=e, countdown=60 * (2 ** self.request.retries))
        
    except OpenAIAPIError as e:
        # Other API errors - retry with exponential backoff
        error_msg = f"OpenAI API error: {str(e)}"
        logger.error(f"API error for job {job_id}: {error_msg}")
        WhisperAPICircuitBreaker.record_failure()
        raise self.retry(exc=e, countdown=60 * (2 ** self.request.retries))
        
    except Exception as e:
        # Other errors - mark job as failed
        error_msg = f"Transcription failed: {str(e)}"
        logger.error(f"Error transcribing job {job_id}: {error_msg}")
        job.status = 'failed'
        job.error_message = error_msg
        job.completed_at = timezone.now()
        job.save()
        WhisperAPICircuitBreaker.record_failure()
        raise
    finally:
        # Clean up input file after processing (whether success or failure)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up input file: {file_path}")
        except Exception as cleanup_error:
            logger.warning(f"Failed to clean up file {file_path}: {cleanup_error}")


@shared_task(
    bind=True,
    soft_time_limit=60,  # 1 minute
    time_limit=120,  # 2 minutes
    queue='lightweight'
)
def generate_docx_task(self, job_id, transcription_text):
    """
    Generate DOCX file from transcription text.
    
    Args:
        job_id: UUID of the TranscriptionJob
        transcription_text: Text to convert to DOCX
    
    Returns:
        Path to generated DOCX file
    """
    try:
        from docx import Document
        import io
        
        job = TranscriptionJob.objects.get(job_id=job_id)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('Audio Transcription', 0)
        
        # Add transcription text
        # Split by paragraphs (double newlines) for better formatting
        if '\n\n' in transcription_text:
            paragraphs = transcription_text.split('\n\n')
        else:
            paragraphs = transcription_text.split('\n')
        
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to temporary file
        temp_dir = tempfile.mkdtemp()
        timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')
        docx_filename = f'transcription-{timestamp}.docx'
        docx_path = os.path.join(temp_dir, docx_filename)
        
        doc.save(docx_path)
        
        logger.info(f"DOCX generated for job {job_id}: {docx_path}")
        
        return docx_path
        
    except Exception as e:
        error_msg = f"Error generating DOCX: {str(e)}"
        logger.error(f"Error generating DOCX for job {job_id}: {error_msg}")
        raise
