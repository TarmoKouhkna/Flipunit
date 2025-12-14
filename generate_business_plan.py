#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Skript FlipUnit.eu äriplaani Wordi dokumendi genereerimiseks
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        tag = 'w:{}'.format(margin_name)
        element = OxmlElement(tag)
        element.set(qn('w:w'), str(int(margin_value * 20)))  # Convert to twips
        element.set(qn('w:type'), 'dxa')
        tcMar.append(element)
    tcPr.append(tcMar)


def add_heading_with_style(doc, text, level=1):
    """Add heading with custom style"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_table_with_data(doc, headers, rows, title=None):
    """Add formatted table with proper column widths and no text wrapping"""
    if title:
        doc.add_paragraph(title, style='Normal')
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    # Use a simpler table style for better readability
    table.style = 'Table Grid'
    
    # Set column widths - make them wider to prevent text wrapping
    # Total width should be around 7 inches (page width minus margins)
    num_cols = len(headers)
    if num_cols == 6:  # Financial tables: Näitaja, Q1, Q2, Q3, Q4, Aasta kokku
        # Much wider columns to prevent wrapping
        col_widths = [Inches(2.8), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.5)]
    else:
        # Default: first column wider, others equal
        first_col_width = Inches(2.5)
        remaining_width = Inches(4.5) / (num_cols - 1) if num_cols > 1 else Inches(1.5)
        col_widths = [first_col_width] + [remaining_width] * (num_cols - 1)
    
    # Apply widths to all rows and set row height
    for row_idx, row in enumerate(table.rows):
        # Set minimum row height for compact display
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '240')  # Minimum height in twips (12pt)
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        for idx, cell in enumerate(row.cells):
            if idx < len(col_widths):
                cell.width = col_widths[idx]
    
    # Disable autofit to keep our widths
    table.autofit = False
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(header_cells):
            # Clear existing text
            header_cells[i].text = ""
            para = header_cells[i].paragraphs[0]
            run = para.add_run(str(header))
            run.font.bold = True
            run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = 1.0  # Single line spacing
            # Disable text wrapping
            para.paragraph_format.widow_control = False
            # Set cell margins - very tight margins
            set_cell_margins(header_cells[i], top=10, bottom=10, left=10, right=10)
            # Set vertical alignment
            header_cells[i].vertical_alignment = 1  # Center
            # Prevent text wrapping in cell
            tcPr = header_cells[i]._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(col_widths[i].inches * 1440)))  # Convert to twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    
    # Add rows
    for row_idx, row_data in enumerate(rows, start=1):
        if row_idx < len(table.rows):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row_cells) and col_idx < len(headers):
                    cell_text = str(cell_data) if cell_data else ""
                    # Clear existing text
                    row_cells[col_idx].text = ""
                    para = row_cells[col_idx].paragraphs[0]
                    run = para.add_run(cell_text)
                    run.font.size = Pt(9)  # Compact font size
                    # Center align for numeric columns (except first)
                    if col_idx > 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.line_spacing = 1.0  # Single line spacing
                    # Prevent text wrapping
                    para.paragraph_format.widow_control = False
                    # Set cell margins - very tight margins
                    set_cell_margins(row_cells[col_idx], top=5, bottom=5, left=10, right=10)
                    # Set vertical alignment
                    row_cells[col_idx].vertical_alignment = 1  # Center
                    # Prevent text wrapping in cell - set cell width explicitly
                    tcPr = row_cells[col_idx]._tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(int(col_widths[col_idx].inches * 1440)))  # Convert to twips
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
    
    doc.add_paragraph()  # Add spacing after table
    return table


def add_bullet_list(doc, items):
    """Add bullet list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()  # Add spacing after list


def create_business_plan_document():
    """Create the complete business plan document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('FlipUnit.eu - Äriplaan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # ========== KOKKUVÕTE ==========
    doc.add_heading('Kokkuvõte', 1)
    doc.add_paragraph(
        'FlipUnit.eu on tasuta veebipõhine konverterite platvorm, mis pakub 100+ tööriista 12 kategoorias '
        'ilma registreerimise või sisselogimiseta. Platvorm on mõeldud kasutajatele, kes vajavad kiireid '
        'konversioone ilma tarkvara paigaldamise või kontode loomise vajaduseta. Platvormi eesmärk on pakkuda '
        'kõike ühes kohas, tagades kiire, turvalise ja privaatse kasutuskogemuse.'
    )
    doc.add_paragraph()
    
    # ========== 1. TOOTE ÜLEVADE ==========
    doc.add_heading('1. Toote ülevaade', 1)
    
    doc.add_heading('1.1 Põhiväärtus', 2)
    doc.add_paragraph('FlipUnit.eu pakub:')
    add_bullet_list(doc, [
        'Ühe koha lahendus kõigile konversioonivajadustele',
        'Registreerimise või sisselogimise vajadus puudub',
        'Kiire ja turvaline töötlemine',
        'Mobiilseadmetele optimeeritud disain',
        'Privaatsus (failid töödeldakse ja kustutatakse automaatselt)'
    ])
    
    doc.add_heading('1.2 Toote kategooriad ja tööriistad', 2)
    
    doc.add_heading('1.2.1 Mõõtühikute konverterid (6 tööriista)', 3)
    add_bullet_list(doc, [
        'Pikkuse konverter: meetrid, kilomeetrid, sentimeetrid, millimeetrid, miilid, jalad, tollid, jardid',
        'Kaalu konverter: kilogrammid, grammid, naelad, untsid, tonnid, stone',
        'Temperatuuri konverter: Celsiuse, Fahrenheiti, Kelvini kraadid',
        'Mahu konverter: liitrid, milliliitrid, gallonid (USA/UK), kvartid, pintid, tassid, vedeliku untsid',
        'Pindala konverter: ruutmeetrid, ruutjalad, aakrid, hektaarid',
        'Kiiruse konverter: km/h, miil/h, m/s, sõlmed'
    ])
    
    doc.add_heading('1.2.2 Pildi konverterid (9 tööriista)', 3)
    doc.add_paragraph('Formaadi konversioon: JPEG ↔ PNG, WebP ↔ PNG/JPG, SVG → PNG')
    doc.add_paragraph('Pildiredigeerimine:')
    add_bullet_list(doc, [
        'Suuruse muutmine (aspect ratio säilitamine)',
        'Pööramine ja peegeldamine',
        'Vesimärgi lisamine',
        'Halltoonideks teisendamine',
        'Piltide ühendamine',
        'EXIF andmete eemaldamine'
    ])
    
    doc.add_heading('1.2.3 Meedia konverterid (9 tööriista)', 3)
    add_bullet_list(doc, [
        'Video/audio konversioon: MP4, AVI, MOV, MKV, WebM',
        'Audio konversioon: MP3, WAV, FLAC, AAC, OGG, M4A',
        'MP4 → MP3: audio eraldamine videost',
        'Video → GIF: videofailide teisendamine animatsiooniks',
        'Audio tööriistad: audio failide jagamine, mitme faili ühendamine, video kompressioon, video heli vaigistamine, taustamüra vähendamine'
    ])
    
    doc.add_heading('1.2.4 PDF tööriistad (11 tööriista)', 3)
    doc.add_paragraph('PDF manipuleerimine:')
    add_bullet_list(doc, [
        'Mitme PDF faili ühendamine',
        'PDF failide jagamine lehtedeks',
        'PDF → pildid (PNG, JPEG)',
        'PDF → HTML',
        'HTML → PDF',
        'PDF → tekst',
        'PDF kompressioon',
        'PDF lehtede pööramine',
        'OCR (skaneeritud PDF-de tekstituvastus)',
        'Metaandmete eemaldamine'
    ])
    
    doc.add_heading('1.2.5 Valuuta ja krüpto konverter', 3)
    add_bullet_list(doc, [
        'Reaalajas valuutakursid',
        'Kõik peamised valuutad (EUR, USD, GBP, JPY, jne)',
        'Krüptovaluutad (BTC, ETH, jne)',
        'Kulla hind (trooja unts, gramm, kilogramm)'
    ])
    
    doc.add_heading('1.2.6 Arhiivi konverterid (7 tööriista)', 3)
    doc.add_paragraph('Arhiiviformaatide konversioon:')
    add_bullet_list(doc, [
        'RAR → ZIP',
        'ZIP → 7Z',
        '7Z → ZIP',
        'TAR.GZ ↔ ZIP',
        'ISO failide ekstraheerimine',
        'ZIP arhiivide loomine'
    ])
    
    doc.add_heading('1.2.7 Teksti ja stringi konverterid (10 tööriista)', 3)
    doc.add_paragraph('Teksti teisendused:')
    add_bullet_list(doc, [
        'Suur- ja väiketähed',
        'CamelCase ↔ snake_case',
        'Eri märkide eemaldamine',
        'Duplikaatridade eemaldamine',
        'Ridade sorteerimine'
    ])
    doc.add_paragraph('Formaadi konversioonid:')
    add_bullet_list(doc, [
        'JSON ↔ XML',
        'JSON ↔ YAML',
        'HTML ↔ Markdown',
        'Tekst ↔ Base64'
    ])
    doc.add_paragraph('Teksti analüüs: sõnade, tähemärkide, ridade loendamine')
    
    doc.add_heading('1.2.8 Arendaja tööriistad (10 tööriista)', 3)
    doc.add_paragraph('Koodi töötlemine:')
    add_bullet_list(doc, [
        'HTML/CSS/JavaScript minifitseerimine',
        'Minifitseeritud koodi ilutamine'
    ])
    doc.add_paragraph('Andmete konversioonid:')
    add_bullet_list(doc, [
        'CSV ↔ JSON',
        'SQL päringute vormindamine',
        'CSS ↔ SCSS'
    ])
    doc.add_paragraph('Arendaja tööriistad:')
    add_bullet_list(doc, [
        'Regulaaravaldiste testija',
        'JWT tokenite dekodeerimine',
        'URL kodeerimine/dekodeerimine',
        'Räsiväärtuste genereerimine (MD5, SHA1, SHA256)'
    ])
    
    doc.add_heading('1.2.9 Utiliidid (12 tööriista)', 3)
    add_bullet_list(doc, [
        'Kalkulaator: põhilised matemaatilised tehted',
        'QR-koodi generaator: teksti/URL-ide põhjal',
        'Ajavööndi konverter: ajavööndite vahel teisendamine',
        'Rooma numbrite konverter: Rooma ↔ araabia numbrid',
        'Faviconi generaator: pildist favicon.ico loomine',
        'Ajatempli konverter: ajatemplid ↔ kuupäevad',
        'Tekst-kõne: teksti teisendamine kõneks',
        'Juhuslike numbrite generaator',
        'Lorem Ipsum generaator',
        'Juhuslike sõnade generaator',
        'Juhuslike nimede generaator',
        'Sõnaloterii'
    ])
    
    doc.add_heading('1.2.10 Värvi valija tööriistad', 3)
    add_bullet_list(doc, [
        'Ekraanivärvi valija: värvide valimine ekraanilt',
        'Pildist värvide eraldamine',
        'Värviformaatide konversioon: HEX, RGB, HSL, CMYK'
    ])
    
    doc.add_heading('1.2.11 Veebisaidi staatuse kontrollija', 3)
    add_bullet_list(doc, [
        'Reaalajas veebisaidi kättesaadavuse kontroll',
        'Vastuse aja mõõtmine',
        'HTTP staatuse koodide analüüs'
    ])
    
    doc.add_heading('1.2.12 YouTube pisipildi allalaadija', 3)
    doc.add_paragraph('Pisipiltide allalaadimine erinevates eraldusvõimetes:')
    add_bullet_list(doc, [
        'HD (1280x720)',
        'HQ (640x480)',
        'MQ (320x180)',
        'Vaikimisi (120x90)'
    ])
    
    doc.add_paragraph('Kokku: 100+ tööriista 12 kategoorias')
    doc.add_paragraph()
    
    # ========== 2. TURU ANALÜÜS ==========
    doc.add_heading('2. Turu analüüs', 1)
    
    doc.add_heading('2.1 Sihtrühmad', 2)
    
    doc.add_heading('2.1.1 Peamised kasutajad', 3)
    doc.add_paragraph('Üldkasutajad:')
    add_bullet_list(doc, [
        'Pildi- ja PDF-konversioonid',
        'Dokumentide töötlemine',
        'Igapäevased konversioonid'
    ])
    doc.add_paragraph('Õpilased ja õpetajad:')
    add_bullet_list(doc, [
        'Mõõtühikute konversioonid',
        'Dokumentide tööriistad',
        'Õppematerjalide töötlemine'
    ])
    doc.add_paragraph('Väikeettevõtjad:')
    add_bullet_list(doc, [
        'Dokumentide töötlemine',
        'Pildiredigeerimine',
        'Meediafailide konversioonid'
    ])
    
    doc.add_heading('2.1.2 Teisene sihtrühm', 3)
    doc.add_paragraph('Veebiarendajad:')
    add_bullet_list(doc, [
        'Koodi tööriistad',
        'Formaadi konversioonid',
        'API integratsioonid'
    ])
    doc.add_paragraph('Sisuloojad:')
    add_bullet_list(doc, [
        'Meedia konversioonid',
        'Pildiredigeerimine',
        'Video töötlemine'
    ])
    doc.add_paragraph('Professionaalid:')
    add_bullet_list(doc, [
        'PDF tööriistad',
        'Valuuta konversioonid',
        'Dokumentide töötlemine'
    ])
    
    doc.add_heading('2.2 Turu suurus', 2)
    doc.add_paragraph('Globaalne turu ülevaade:')
    add_bullet_list(doc, [
        'Veebipõhiste tööriistade turg: hinnanguliselt 5–10 miljardit eurot (kasv 15–20% aastas)',
        'Konverterite turg: kõrge otsingumaht (miljonid päevas)',
        'Registreerimiseta tööriistad: kasvav nõudlus privaatsuse tõttu'
    ])
    
    doc.add_paragraph('Eesti turu potentsiaal:')
    add_bullet_list(doc, [
        'Eesti elanikkond: ~1,3 miljonit',
        'Internetikasutajad: ~95% (1,2 miljonit)',
        'Potentsiaalne kasutajabaas: 50,000–100,000 aktiivset kasutajat'
    ])
    
    doc.add_paragraph('Rahvusvaheline turu potentsiaal:')
    add_bullet_list(doc, [
        'Euroopa: 500+ miljonit internetikasutajat',
        'Põhjusmärksõnad: miljonid päevas',
        'Potentsiaalne kasutajabaas: 1–5 miljonit kuus'
    ])
    
    doc.add_heading('2.3 Konkurentsianalüüs', 2)
    
    doc.add_heading('2.3.1 Peamised konkurendid', 3)
    doc.add_paragraph('1. Zamzar')
    add_bullet_list(doc, [
        'Tugevused: lai valik, API',
        'Nõrkused: registreerimine, aeglane töötlemine'
    ])
    doc.add_paragraph('2. CloudConvert')
    add_bullet_list(doc, [
        'Tugevused: API, pilve integratsioonid',
        'Nõrkused: keerukas kasutajaliides, kõrged hinnad'
    ])
    doc.add_paragraph('3. Convertio')
    add_bullet_list(doc, [
        'Tugevused: lihtne kasutajaliides',
        'Nõrkused: piiratud valik, reklaamid'
    ])
    doc.add_paragraph('4. Online-Convert')
    add_bullet_list(doc, [
        'Tugevused: lai valik',
        'Nõrkused: registreerimine, aeglane töötlemine'
    ])
    
    doc.add_heading('2.3.2 Konkurentsieelised', 3)
    add_bullet_list(doc, [
        'Laiem tööriistade valik (100+ vs 20–50)',
        'Registreerimise vajadus puudub',
        'Kiirem töötlemine',
        'Parem kasutajakogemus',
        'Privaatsus (failid kustutatakse automaatselt)',
        'Mobiilne optimeerimine',
        'SEO optimeerimine'
    ])
    
    # ========== 3. ÄRIMUDEL ==========
    doc.add_heading('3. Ärimudel', 1)
    
    doc.add_heading('3.1 Praegune mudel: Freemium (tasuta tasand)', 2)
    doc.add_paragraph('Põhifunktsioonid:')
    add_bullet_list(doc, [
        'Registreerimise vajadus puudub',
        'Tasuta juurdepääs kõigile tööriistadele',
        'Kiire töötlemine (kuni 700MB failid)',
        'Privaatsus (ajutine failide salvestamine)',
        'Põhiliselt reklaamipõhine'
    ])
    
    doc.add_heading('3.2 Tulevased tasustamise võimalused', 2)
    
    doc.add_heading('3.2.1 Premium tellimus (eelistatud)', 3)
    doc.add_paragraph('Premium funktsioonid:')
    add_bullet_list(doc, [
        'Suuremad failid (kuni 2GB+)',
        'Partsi töötlemine (mitme faili korraga)',
        'Prioriteetne töötlemine',
        'API juurdepääs',
        'Ilma reklaamideta kogemus',
        'Pilve salvestusruumi integratsioon',
        'Ajaloo säilitamine',
        'Eraldi toetusteenus'
    ])
    doc.add_paragraph('Hinnakujundus:')
    add_bullet_list(doc, [
        'Kuu: 4,99–9,99 €',
        'Aasta: 49,99–99,99 € (2 kuud tasuta)'
    ])
    
    doc.add_heading('3.2.2 Reklaamitulu', 3)
    add_bullet_list(doc, [
        'Google AdSense',
        'Media.net',
        'Sponsooritud tööriistade paigutused',
        'Partnerreklaamid'
    ])
    doc.add_paragraph('Eeldatav tulu:')
    add_bullet_list(doc, [
        'CPM: 2–5 €',
        'CTR: 1–3%',
        'Kuu: 500–2,500 € (olenevalt liiklusest)'
    ])
    
    doc.add_heading('3.2.3 Ettevõtete/B2B lahendused', 3)
    add_bullet_list(doc, [
        'Valge märgistus',
        'API juurdepääs ettevõtetele',
        'Kohandatud integratsioonid',
        'Eraldi toetusteenus',
        'SLA garantiid'
    ])
    doc.add_paragraph('Hinnakujundus:')
    add_bullet_list(doc, [
        'Starter: 29,99 €/kuu (10,000 API päringut)',
        'Professional: 99,99 €/kuu (100,000 API päringut)',
        'Enterprise: kohandatud (piiramatud päringud)'
    ])
    
    doc.add_heading('3.2.4 Tasu kasutamise mudel', 3)
    add_bullet_list(doc, [
        'Krediidid premium konversioonidele',
        'Ühekordne tasumine suurte failide eest',
        'Krediidipaketid: 5 € (50 krediiti), 10 € (120 krediiti), 20 € (250 krediiti)'
    ])
    
    # ========== 4. TULUDE ALLIKAD ==========
    doc.add_heading('4. Tulude allikad', 1)
    
    doc.add_heading('4.1 Esimese aasta prognoosid', 2)
    
    doc.add_heading('4.1.1 Reklaamitulu (peamine)', 3)
    doc.add_paragraph('Kuu 1–3:')
    add_bullet_list(doc, [
        'Hinnanguline liiklus: 5,000–10,000 külastajat/kuu',
        'CPM: 2–3 €',
        'Hinnanguline tulu: 200–600 €/kuu'
    ])
    doc.add_paragraph('Kuu 4–6:')
    add_bullet_list(doc, [
        'Hinnanguline liiklus: 15,000–30,000 külastajat/kuu',
        'CPM: 2,5–3,5 €',
        'Hinnanguline tulu: 600–1,500 €/kuu'
    ])
    doc.add_paragraph('Kuu 7–9:')
    add_bullet_list(doc, [
        'Hinnanguline liiklus: 30,000–50,000 külastajat/kuu',
        'CPM: 3–4 €',
        'Hinnanguline tulu: 1,200–2,500 €/kuu'
    ])
    doc.add_paragraph('Kuu 10–12:')
    add_bullet_list(doc, [
        'Hinnanguline liiklus: 50,000–100,000 külastajat/kuu',
        'CPM: 3,5–5 €',
        'Hinnanguline tulu: 2,000–5,000 €/kuu'
    ])
    doc.add_paragraph('Aasta kokku: 12,000–36,000 €')
    
    doc.add_heading('4.1.2 Premium tellimused (teisene)', 3)
    doc.add_paragraph('Kuu 1–3:')
    add_bullet_list(doc, [
        'Tellimuste arv: 0–10',
        'Keskmine hind: 6,99 €/kuu',
        'Tulu: 0–70 €/kuu'
    ])
    doc.add_paragraph('Kuu 4–6:')
    add_bullet_list(doc, [
        'Tellimuste arv: 25–50',
        'Keskmine hind: 6,99 €/kuu',
        'Tulu: 175–350 €/kuu'
    ])
    doc.add_paragraph('Kuu 7–9:')
    add_bullet_list(doc, [
        'Tellimuste arv: 50–100',
        'Keskmine hind: 6,99 €/kuu',
        'Tulu: 350–700 €/kuu'
    ])
    doc.add_paragraph('Kuu 10–12:')
    add_bullet_list(doc, [
        'Tellimuste arv: 100–200',
        'Keskmine hind: 6,99 €/kuu',
        'Tulu: 700–1,400 €/kuu'
    ])
    doc.add_paragraph('Aasta kokku: 3,600–7,200 €')
    
    doc.add_heading('4.1.3 Partnerreklaamid', 3)
    add_bullet_list(doc, [
        'Tarkvara/teenuste soovitused',
        'Kuu: 100–500 €',
        'Aasta: 1,200–6,000 €'
    ])
    
    doc.add_heading('4.1.4 B2B lahendused (aasta lõpus)', 3)
    add_bullet_list(doc, [
        '2–5 ettevõtte klienti',
        'Keskmine: 50 €/kuu',
        'Aasta: 1,200–3,000 €'
    ])
    
    doc.add_paragraph('Esimese aasta kogutulu: 18,000–52,200 €')
    
    doc.add_heading('4.2 Teise ja kolmanda aasta prognoosid', 2)
    doc.add_paragraph('Teine aasta:')
    add_bullet_list(doc, [
        'Liiklus: 200,000–500,000 külastajat/kuu',
        'Reklaamitulu: 5,000–15,000 €/kuu',
        'Premium tellimused: 2,000–5,000 €/kuu',
        'B2B: 2,000–5,000 €/kuu',
        'Kogutulu: 108,000–300,000 €/aasta'
    ])
    
    doc.add_paragraph('Kolmas aasta:')
    add_bullet_list(doc, [
        'Liiklus: 500,000–1,000,000 külastajat/kuu',
        'Reklaamitulu: 15,000–40,000 €/kuu',
        'Premium tellimused: 5,000–15,000 €/kuu',
        'B2B: 5,000–15,000 €/kuu',
        'Kogutulu: 300,000–840,000 €/aasta'
    ])
    
    # ========== 5. TURUNDUSSTRATEEGIA ==========
    doc.add_heading('5. Turundusstrateegia', 1)
    
    doc.add_heading('5.1 SEO ja sisuturundus', 2)
    doc.add_paragraph('Praegune SEO:')
    add_bullet_list(doc, [
        'XML sitemap',
        'Meta sildid ja Open Graph',
        'Mobiilne optimeerimine',
        'Kiire laadimisaeg',
        'Struktureeritud andmed'
    ])
    
    doc.add_paragraph('Tulevased SEO meetmed:')
    add_bullet_list(doc, [
        'Blogi: tööriistade juhendid',
        'Põhjusmärksõnade sihtimine',
        'Sisemine linkimine',
        'Kohalik SEO (Eesti)',
        'Mitmekeelne sisu (eesti, inglise, vene)'
    ])
    
    doc.add_paragraph('Eeldatav tulemus:')
    add_bullet_list(doc, [
        '6–12 kuud: esimesed 10 tulemust',
        '12–24 kuud: esimesed 3 tulemust valitud märksõnadele',
        'Otsinguliiklus: 50,000–200,000 külastajat/kuu'
    ])
    
    doc.add_heading('5.2 Sotsiaalmeedia turundus', 2)
    doc.add_paragraph('YouTube:')
    add_bullet_list(doc, [
        'Tööriistade demo videod',
        '"Kuidas" juhendid',
        'Eeldatav: 1,000–5,000 tellijat esimesel aastal'
    ])
    
    doc.add_paragraph('TikTok/Instagram Reels:')
    add_bullet_list(doc, [
        'Lühikesed demo videod',
        'Trendide jälgimine',
        'Eeldatav: 5,000–20,000 jälgijat'
    ])
    
    doc.add_paragraph('Twitter/X:')
    add_bullet_list(doc, [
        'Uuendused ja näpunäited',
        'Kogukonna interaktsioon',
        'Eeldatav: 500–2,000 jälgijat'
    ])
    
    doc.add_paragraph('Reddit:')
    add_bullet_list(doc, [
        'Kogukondade osalemine',
        'Kasulikud postitused',
        'Eeldatav: 1,000–5,000 liiget'
    ])
    
    doc.add_paragraph('Product Hunt:')
    add_bullet_list(doc, [
        'Platvormi käivitamine',
        'Eeldatav: 100–500 häält'
    ])
    
    doc.add_heading('5.3 Partnerlused', 2)
    doc.add_paragraph('Integratsioonid:')
    add_bullet_list(doc, [
        'Tootlikkustööriistad (Notion, Trello)',
        'Pilve salvestusruum (Google Drive, Dropbox)',
        'Arendajate tööriistad (GitHub, GitLab)'
    ])
    
    doc.add_paragraph('Kogukonna partnerlused:')
    add_bullet_list(doc, [
        'Arendajate kogukonnad',
        'Ülikoolid ja koolid',
        'Väikeettevõtted'
    ])
    
    doc.add_paragraph('Eeldatav tulemus:')
    add_bullet_list(doc, [
        '10–20 partnerit esimesel aastal',
        '50–100 partnerit teisel aastal'
    ])
    
    doc.add_heading('5.4 Tasustatud reklaam', 2)
    doc.add_paragraph('Google Ads:')
    add_bullet_list(doc, [
        'Kõrge kavatsusega märksõnad',
        'Eeldatav: 2,000–5,000 €/kuu',
        'ROI: 3–5x'
    ])
    
    doc.add_paragraph('Facebook/Instagram reklaam:')
    add_bullet_list(doc, [
        'Sihtrühmade reklaamid',
        'Eeldatav: 1,000–3,000 €/kuu',
        'ROI: 2–4x'
    ])
    
    doc.add_paragraph('Retargeting:')
    add_bullet_list(doc, [
        'Tagasipöördumise kampaaniad',
        'Eeldatav: 500–1,500 €/kuu',
        'ROI: 4–6x'
    ])
    
    doc.add_paragraph('Kokku: 3,500–9,500 €/kuu')
    
    # ========== 6. TEHNOLOOGIA JA OPERATSIOONID ==========
    doc.add_heading('6. Tehnoloogia ja operatsioonid', 1)
    
    doc.add_heading('6.1 Tehnoloogia stack', 2)
    doc.add_paragraph('Backend:')
    add_bullet_list(doc, [
        'Django 5.2.8 (Python 3.12)',
        'PostgreSQL andmebaas',
        'Gunicorn WSGI server',
        'WhiteNoise staatiliste failide teenindamiseks'
    ])
    
    doc.add_paragraph('Töötlemise teegid:')
    add_bullet_list(doc, [
        'Pildid: Pillow 12.0.0, CairoSVG 2.8.2',
        'PDF: pypdf 6.3.0, pdf2image 1.17.0, WeasyPrint 62.3',
        'Meedia: FFmpeg (süsteemi sõltuvus)',
        'Arhiivid: py7zr, rarfile, pycdlib',
        'Tekst: PyYAML, markdown, html2text'
    ])
    
    doc.add_paragraph('Deploy:')
    add_bullet_list(doc, [
        'Docker ja Docker Compose',
        'Nginx reverse proxy',
        'SSL/TLS sertifikaadid (Let\'s Encrypt)'
    ])
    
    doc.add_paragraph('Hosting:')
    add_bullet_list(doc, [
        'VPS/Cloud (hetkel juurutatud)',
        'Tulevane: pilve infrastruktuur (AWS, Google Cloud, Azure)'
    ])
    
    doc.add_heading('6.2 Infrastruktuuri kulud', 2)
    doc.add_paragraph('Praegused kulud:')
    add_bullet_list(doc, [
        'Server hosting: 50–200 €/kuu',
        'CDN: 20–100 €/kuu',
        'Andmebaas: 30–100 €/kuu (või kaasas)',
        'Domeen ja SSL: 20 €/aasta',
        'Kokku: 100–400 €/kuu'
    ])
    
    doc.add_paragraph('Skaleerimise kulud (2. aasta):')
    add_bullet_list(doc, [
        'Server hosting: 200–500 €/kuu',
        'CDN: 100–300 €/kuu',
        'Andmebaas: 100–300 €/kuu',
        'Töötlemise serverid: 100–200 €/kuu',
        'Kokku: 500–1,300 €/kuu'
    ])
    
    doc.add_paragraph('Skaleerimise kulud (3. aasta):')
    add_bullet_list(doc, [
        'Server hosting: 500–1,000 €/kuu',
        'CDN: 300–800 €/kuu',
        'Andmebaas: 300–800 €/kuu',
        'Töötlemise serverid: 200–500 €/kuu',
        'Kokku: 1,300–3,100 €/kuu'
    ])
    
    doc.add_heading('6.3 Operatiivne tõhusus', 2)
    doc.add_paragraph('Automaatne failide puhastus:')
    add_bullet_list(doc, [
        'Failid kustutatakse automaatselt 1 tunni pärast',
        'Andmebaasi puhastus: iga päev',
        'Logide puhastus: iga nädal'
    ])
    
    doc.add_paragraph('Tõhus töötlemise algoritm:')
    add_bullet_list(doc, [
        'Optimeeritud kood',
        'Vahemälu',
        'Asünkroonne töötlemine (tulevikus)'
    ])
    
    doc.add_paragraph('Skaleeritav arhitektuur:')
    add_bullet_list(doc, [
        'Mikroteenuste arhitektuur (tulevikus)',
        'Laadi jaotamine',
        'Auto-skaleerimine'
    ])
    
    doc.add_paragraph('Pilve salvestusruumi integratsioon (tulevikus):')
    add_bullet_list(doc, [
        'AWS S3 või Google Cloud Storage',
        'Failide CDN-iga levitamine',
        'Kulud: 50–200 €/kuu'
    ])
    
    # ========== 7. KONKURENTSIEELISED ==========
    doc.add_heading('7. Konkurentsieelised', 1)
    
    doc.add_heading('7.1 Toote eelised', 2)
    add_bullet_list(doc, [
        'Lai tööriistade valik: 100+ tööriista vs 20–50 konkurentidel',
        'Registreerimise vajadus puudub: väiksem takistus',
        'Kiire töötlemine: optimeeritud algoritmid',
        'Privaatsus: ajutine failide salvestamine',
        'Mobiilne optimeerimine: kõik seadmed',
        'Arendajatele sobiv: puhtas API-valmis arhitektuur',
        'SEO optimeerimine: tugev otsingunähtavus'
    ])
    
    doc.add_heading('7.2 Tehnoloogilised eelised', 2)
    add_bullet_list(doc, [
        'Kaasaegne tehnoloogia stack',
        'Skaleeritav arhitektuur',
        'Kiire laadimisaeg',
        'Turvalisus (Django turvalisus)',
        'Avatud lähtekood (võimalus)'
    ])
    
    doc.add_heading('7.3 Kasutajakogemuse eelised', 2)
    add_bullet_list(doc, [
        'Lihtne kasutajaliides',
        'Kiire töötlemine',
        'Selged veateated',
        'Responsiivne disain',
        'Kohandatud kogemus'
    ])
    
    # ========== 8. KASVUSTRATEEGIA ==========
    doc.add_heading('8. Kasvustrateegia', 1)
    
    doc.add_heading('8.1 Faas 1: Alus (kuud 1–6)', 2)
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        'Platvormi käivitamine ja stabiliseerimine',
        'SEO optimeerimine',
        'Kasutajate tagasiside kogumine',
        'Jõudluse optimeerimine',
        'Põhiline turundus'
    ])
    
    doc.add_paragraph('Meetmed:')
    add_bullet_list(doc, [
        'Google Search Console seadistamine',
        'Google Analytics integratsioon',
        'Tagasiside süsteem',
        'Jõudluse jälgimine',
        'Põhiline reklaam'
    ])
    
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        '10,000–25,000 külastajat/kuu',
        '1,000–5,000 aktiivset kasutajat/kuu',
        '50–100 tagasisidet/kuu'
    ])
    
    doc.add_paragraph('Kulud: 1,000–2,000 €/kuu')
    
    doc.add_heading('8.2 Faas 2: Kasv (kuud 7–12)', 2)
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        'Premium funktsioonide juurutamine',
        'Reklaamitulu käivitamine',
        'Sisuturundus',
        'Sotsiaalmeedia kohalolek',
        'Partnerlused'
    ])
    
    doc.add_paragraph('Meetmed:')
    add_bullet_list(doc, [
        'Premium tellimuste süsteem',
        'Google AdSense integratsioon',
        'Blogi käivitamine',
        'Sotsiaalmeedia kontod',
        'Partnerite otsimine'
    ])
    
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        '50,000–100,000 külastajat/kuu',
        '5,000–10,000 aktiivset kasutajat/kuu',
        '100–200 premium tellimust',
        '2,000–5,000 € tulu/kuu'
    ])
    
    doc.add_paragraph('Kulud: 2,000–5,000 €/kuu')
    
    doc.add_heading('8.3 Faas 3: Skaleerimine (2. aasta)', 2)
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        'API arendamine',
        'Ettevõtete lahendused',
        'Mobiilirakendus (valikuline)',
        'Rahvusvaheline laienemine',
        'Täiendavad funktsioonid'
    ])
    
    doc.add_paragraph('Meetmed:')
    add_bullet_list(doc, [
        'RESTful API arendamine',
        'B2B müügitiim',
        'Mobiilirakenduse arendamine',
        'Mitmekeelne tugi',
        'Uute tööriistade lisamine'
    ])
    
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        '200,000–500,000 külastajat/kuu',
        '20,000–50,000 aktiivset kasutajat/kuu',
        '500–1,000 premium tellimust',
        '10–20 B2B klienti',
        '10,000–25,000 € tulu/kuu'
    ])
    
    doc.add_paragraph('Kulud: 5,000–10,000 €/kuu')
    
    doc.add_heading('8.4 Faas 4: Tasustamine (2.–3. aasta)', 2)
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        'Premium tellimuste käivitamine',
        'Ettevõtete partnerlused',
        'Valge märgistus',
        'Täiendavad tulude allikad'
    ])
    
    doc.add_paragraph('Meetmed:')
    add_bullet_list(doc, [
        'Premium tellimuste turundus',
        'B2B müügikampaaniad',
        'Valge märgistuse pakkumine',
        'Uute tulude allikate uurimine'
    ])
    
    doc.add_paragraph('Eesmärgid:')
    add_bullet_list(doc, [
        '500,000–1,000,000 külastajat/kuu',
        '50,000–100,000 aktiivset kasutajat/kuu',
        '1,000–5,000 premium tellimust',
        '50–100 B2B klienti',
        '25,000–70,000 € tulu/kuu'
    ])
    
    doc.add_paragraph('Kulud: 10,000–20,000 €/kuu')
    
    # ========== 9. FINANTSPROGNOOSID ==========
    doc.add_heading('9. Finantsprognoosid', 1)
    
    doc.add_heading('9.1 Esimese aasta finantsprognoos', 2)
    
    headers_q1 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q1 = [
        ["Reklaamitulu", "600 €", "1,800 €", "3,600 €", "6,000 €", "12,000 €"],
        ["Premium tellimused", "0 €", "525 €", "1,050 €", "2,100 €", "3,675 €"],
        ["Partnerreklaamid", "300 €", "400 €", "500 €", "600 €", "1,800 €"],
        ["Kogutulu", "900 €", "2,725 €", "5,150 €", "8,700 €", "17,475 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "600 €", "750 €", "900 €", "1,050 €", "3,300 €"],
        ["Turundus", "1,500 €", "3,000 €", "4,500 €", "6,000 €", "15,000 €"],
        ["Arendus", "1,000 €", "1,500 €", "2,000 €", "2,500 €", "7,000 €"],
        ["Personal (osaline)", "0 €", "0 €", "500 €", "1,000 €", "1,500 €"],
        ["Muud kulud", "200 €", "300 €", "400 €", "500 €", "1,400 €"],
        ["Kogukulud", "3,300 €", "5,550 €", "8,300 €", "11,050 €", "28,200 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "-2,400 €", "-2,825 €", "-3,150 €", "-2,350 €", "-10,725 €"]
    ]
    
    add_table_with_data(doc, headers_q1, rows_q1)
    doc.add_paragraph('Tasuvuspunkt: kuu 14–16')
    
    doc.add_heading('9.2 Teise aasta finantsprognoos', 2)
    
    headers_q2 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q2 = [
        ["Reklaamitulu", "15,000 €", "18,000 €", "21,000 €", "24,000 €", "78,000 €"],
        ["Premium tellimused", "6,000 €", "7,500 €", "9,000 €", "10,500 €", "33,000 €"],
        ["B2B lahendused", "3,000 €", "4,500 €", "6,000 €", "7,500 €", "21,000 €"],
        ["Partnerreklaamid", "800 €", "1,000 €", "1,200 €", "1,400 €", "4,400 €"],
        ["Kogutulu", "24,800 €", "31,000 €", "37,200 €", "43,400 €", "136,400 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "1,200 €", "1,400 €", "1,600 €", "1,800 €", "6,000 €"],
        ["Turundus", "7,500 €", "9,000 €", "10,500 €", "12,000 €", "39,000 €"],
        ["Arendus", "3,000 €", "3,500 €", "4,000 €", "4,500 €", "15,000 €"],
        ["Personal", "2,000 €", "2,500 €", "3,000 €", "3,500 €", "11,000 €"],
        ["Muud kulud", "500 €", "600 €", "700 €", "800 €", "2,600 €"],
        ["Kogukulud", "14,200 €", "17,000 €", "19,800 €", "22,600 €", "73,600 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "10,600 €", "14,000 €", "17,400 €", "20,800 €", "62,800 €"]
    ]
    
    add_table_with_data(doc, headers_q2, rows_q2)
    doc.add_paragraph('Kasumimarginaal: 46%')
    
    doc.add_heading('9.3 Kolmanda aasta finantsprognoos', 2)
    
    headers_q3 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q3 = [
        ["Reklaamitulu", "45,000 €", "50,000 €", "55,000 €", "60,000 €", "210,000 €"],
        ["Premium tellimused", "18,000 €", "21,000 €", "24,000 €", "27,000 €", "90,000 €"],
        ["B2B lahendused", "12,000 €", "15,000 €", "18,000 €", "21,000 €", "66,000 €"],
        ["Partnerreklaamid", "2,000 €", "2,500 €", "3,000 €", "3,500 €", "11,000 €"],
        ["Kogutulu", "77,000 €", "88,500 €", "100,000 €", "111,500 €", "377,000 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "2,400 €", "2,700 €", "3,000 €", "3,300 €", "11,400 €"],
        ["Turundus", "15,000 €", "18,000 €", "21,000 €", "24,000 €", "78,000 €"],
        ["Arendus", "6,000 €", "7,000 €", "8,000 €", "9,000 €", "30,000 €"],
        ["Personal", "5,000 €", "6,000 €", "7,000 €", "8,000 €", "26,000 €"],
        ["Muud kulud", "1,000 €", "1,200 €", "1,400 €", "1,600 €", "5,200 €"],
        ["Kogukulud", "29,400 €", "34,900 €", "40,400 €", "45,900 €", "150,600 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "47,600 €", "53,600 €", "59,600 €", "65,600 €", "226,400 €"]
    ]
    
    add_table_with_data(doc, headers_q3, rows_q3)
    doc.add_paragraph('Kasumimarginaal: 60%')
    
    doc.add_heading('9.4 Investeeringu nõuded', 2)
    doc.add_paragraph('Esimese aasta investeering:')
    add_bullet_list(doc, [
        'Arendus: 7,000 € (juba tehtud)',
        'Turundus: 15,000 €',
        'Infrastruktuur: 3,300 €',
        'Personal: 1,500 €',
        'Muud kulud: 1,400 €',
        'Kokku: 28,200 €'
    ])
    
    doc.add_paragraph('Teise aasta investeering:')
    add_bullet_list(doc, [
        'Arendus: 15,000 €',
        'Turundus: 39,000 €',
        'Infrastruktuur: 6,000 €',
        'Personal: 11,000 €',
        'Muud kulud: 2,600 €',
        'Kokku: 73,600 €'
    ])
    
    doc.add_paragraph('Kolmanda aasta investeering:')
    add_bullet_list(doc, [
        'Arendus: 30,000 €',
        'Turundus: 78,000 €',
        'Infrastruktuur: 11,400 €',
        'Personal: 26,000 €',
        'Muud kulud: 5,200 €',
        'Kokku: 150,600 €'
    ])
    
    doc.add_paragraph('Oodatav ROI:')
    add_bullet_list(doc, [
        '1. aasta: -10,725 € (investeering)',
        '2. aasta: +62,800 € (ROI: 223%)',
        '3. aasta: +226,400 € (ROI: 150%)',
        '3 aasta kogukasum: 278,475 €'
    ])
    
    # ========== 10. RISKIANALÜÜS ==========
    doc.add_heading('10. Riskianalüüs ja leevendamine', 1)
    
    doc.add_heading('10.1 Peamised riskid', 2)
    
    doc.add_heading('10.1.1 Kõrged infrastruktuuri kulud', 3)
    doc.add_paragraph('Risk: Kasutajate kasv suurendab serverikulusid')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine:')
    add_bullet_list(doc, [
        'Tõhus töötlemise algoritm',
        'Pilve auto-skaleerimine',
        'CDN kasutamine',
        'Failide automaatne kustutamine'
    ])
    
    doc.add_heading('10.1.2 Tugev konkurents', 3)
    doc.add_paragraph('Risk: Suuremad konkurendid võivad konkureerida')
    doc.add_paragraph('Tõenäosus: Kõrge')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine:')
    add_bullet_list(doc, [
        'Ainulaadsed funktsioonid',
        'Parem kasutajakogemus',
        'Kiire arendus',
        'Tugev bränd'
    ])
    
    doc.add_heading('10.1.3 Juridilised/järgimisnõuded', 3)
    doc.add_paragraph('Risk: GDPR, andmekaitse, autoriõigused')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Kõrge')
    doc.add_paragraph('Leevendamine:')
    add_bullet_list(doc, [
        'Privaatsuskeskne disain',
        'Selged kasutustingimused',
        'Juridiline nõustamine',
        'Regulaarne audit'
    ])
    
    doc.add_heading('10.1.4 Tehnilised probleemid', 3)
    doc.add_paragraph('Risk: Serveri rikked, turvalisuse rikked')
    doc.add_paragraph('Tõenäosus: Madal')
    doc.add_paragraph('Mõju: Kõrge')
    doc.add_paragraph('Leevendamine:')
    add_bullet_list(doc, [
        'Tugev arhitektuur',
        'Regulaarne jälgimine',
        'Varukoopiad',
        'Kiire taastamine'
    ])
    
    doc.add_heading('10.1.5 Turunduse ebaõnnestumine', 3)
    doc.add_paragraph('Risk: Liikluse kasv on aeglasem kui oodatud')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine:')
    add_bullet_list(doc, [
        'Mitmekesine turundusstrateegia',
        'SEO fookus',
        'Sotsiaalmeedia kohalolek',
        'Partnerlused'
    ])
    
    doc.add_heading('10.2 Riskide haldus', 2)
    doc.add_paragraph('Riskide jälgimine:')
    add_bullet_list(doc, [
        'Igakuine riskide ülevaade',
        'Võtmenäitajate jälgimine',
        'Stsenaariumide planeerimine'
    ])
    
    doc.add_paragraph('Riskide leevendamine:')
    add_bullet_list(doc, [
        'Varufondid',
        'Mitmekesine tulude baas',
        'Tugev tehnoloogia alus',
        'Kogenud meeskond'
    ])
    
    # ========== 11. EDU NÄITAJAD ==========
    doc.add_heading('11. Edu näitajad (KPI-d)', 1)
    
    doc.add_heading('11.1 Kasutaja näitajad', 2)
    doc.add_paragraph('Külastajad:')
    add_bullet_list(doc, [
        'Kuu aktiivsed kasutajad (MAU)',
        'Päev aktiivsed kasutajad (DAU)',
        'Uued kasutajad/kuu',
        'Tagasipöördumise määr'
    ])
    
    doc.add_paragraph('Kasutus:')
    add_bullet_list(doc, [
        'Tööriistade kasutamise sagedus',
        'Keskmine seansi kestus',
        'Lehtede vaatamised seansi kohta',
        'Konversioonide arv kasutaja kohta'
    ])
    
    doc.add_paragraph('Tagasipöördumine:')
    add_bullet_list(doc, [
        '7-päevane tagasipöördumise määr',
        '30-päevane tagasipöördumise määr',
        'Kasutaja elutsükkel'
    ])
    
    doc.add_heading('11.2 Äri näitajad', 2)
    doc.add_paragraph('Tulud:')
    add_bullet_list(doc, [
        'Kuu korduv tulu (MRR)',
        'Aasta korduv tulu (ARR)',
        'Keskmine tellimuse väärtus (ARPU)',
        'Elutsükli väärtus (LTV)'
    ])
    
    doc.add_paragraph('Kulud:')
    add_bullet_list(doc, [
        'Kliendi omandamise maksumus (CAC)',
        'CAC/LTV suhe',
        'Tulu/kasutaja (R/U)'
    ])
    
    doc.add_paragraph('Konversioon:')
    add_bullet_list(doc, [
        'Tasuta → Premium konversioonimäär',
        'Tellimuse konversioonimäär',
        'Reklaamiklikkuste määr (CTR)',
        'Reklaamiklikkuste maksumus (CPC)'
    ])
    
    doc.add_heading('11.3 Tehnilised näitajad', 2)
    doc.add_paragraph('Jõudlus:')
    add_bullet_list(doc, [
        'Lehe laadimisaeg',
        'API vastuse aeg',
        'Faili töötlemise aeg',
        'Töötlemise õnnestumise määr'
    ])
    
    doc.add_paragraph('Kättesaadavus:')
    add_bullet_list(doc, [
        'Ülesaeg protsent',
        'Keskmine rikete aeg (MTBF)',
        'Keskmine taastumise aeg (MTTR)'
    ])
    
    doc.add_paragraph('Turvalisus:')
    add_bullet_list(doc, [
        'Turvalisuse rikete arv',
        'Andmelekete arv',
        'Turvalisuse auditid'
    ])
    
    doc.add_heading('11.4 Eesmärgid', 2)
    doc.add_paragraph('1. aasta:')
    add_bullet_list(doc, [
        'MAU: 10,000–50,000',
        'Konversioonimäär: 1–3%',
        'MRR: 500–2,000 €',
        'CAC: 5–15 €'
    ])
    
    doc.add_paragraph('2. aasta:')
    add_bullet_list(doc, [
        'MAU: 200,000–500,000',
        'Konversioonimäär: 2–5%',
        'MRR: 5,000–15,000 €',
        'CAC: 3–10 €'
    ])
    
    doc.add_paragraph('3. aasta:')
    add_bullet_list(doc, [
        'MAU: 500,000–1,000,000',
        'Konversioonimäär: 3–7%',
        'MRR: 15,000–50,000 €',
        'CAC: 2–8 €'
    ])
    
    # ========== 12. MEESKOND ==========
    doc.add_heading('12. Meeskond ja organisatsioon', 1)
    
    doc.add_heading('12.1 Praegune meeskond', 2)
    doc.add_paragraph('Arendus:')
    add_bullet_list(doc, [
        'Full-stack arendaja (Django, Python)',
        'Frontend arendaja (HTML, CSS, JavaScript)',
        'DevOps insener (Docker, Nginx, serverid)'
    ])
    
    doc.add_paragraph('Turundus:')
    add_bullet_list(doc, [
        'SEO spetsialist',
        'Sotsiaalmeedia spetsialist',
        'Sisuturunduse spetsialist'
    ])
    
    doc.add_paragraph('Toetus:')
    add_bullet_list(doc, [
        'Klienditugi spetsialist (osaline)'
    ])
    
    doc.add_heading('12.2 Tulevased vajadused', 2)
    doc.add_paragraph('1. aasta:')
    add_bullet_list(doc, [
        'Täiendav arendaja (osaline)',
        'Turunduse spetsialist (osaline)'
    ])
    
    doc.add_paragraph('2. aasta:')
    add_bullet_list(doc, [
        'Täisajaline arendaja',
        'Täisajaline turunduse spetsialist',
        'B2B müügispetsialist',
        'Klienditugi spetsialist'
    ])
    
    doc.add_paragraph('3. aasta:')
    add_bullet_list(doc, [
        'Arendusmeeskond (3–5 inimest)',
        'Turundusmeeskond (2–3 inimest)',
        'Müügimeeskond (2–3 inimest)',
        'Toetusmeeskond (2–3 inimest)',
        'Juhtimine (CEO, CTO, CMO)'
    ])
    
    doc.add_heading('12.3 Organisatsiooniline struktuur', 2)
    doc.add_paragraph('Praegune:')
    add_bullet_list(doc, [
        'Tasane struktuur',
        'Väike meeskond',
        'Otsene suhtlus'
    ])
    
    doc.add_paragraph('Tulevane (2.–3. aasta):')
    add_bullet_list(doc, [
        'Funktsionaalne struktuur',
        'Osakonnad: arendus, turundus, müük, toetus',
        'Juhtimise meeskond'
    ])
    
    # ========== 13. JÄRELDUS ==========
    doc.add_heading('13. Järeldus', 1)
    
    doc.add_paragraph(
        'FlipUnit.eu on positsioneeritud kasutama veebipõhiste tööriistade turu kasvu. '
        '100+ tööriistaga, registreerimiseta mudeliga ja skaleeritava arhitektuuriga on platvorm '
        'valmis kasvama tasuta utiliidist tasuliseks platvormiks reklaamitulu, premium tellimuste '
        'ja ettevõtete lahenduste kaudu.'
    )
    
    doc.add_heading('13.1 Peamised edu tegurid', 2)
    add_bullet_list(doc, [
        'Tugev SEO alus',
        'Kasutajakeskne disain',
        'Pidev funktsioonide arendus',
        'Tõhus tasustamisstrateegia',
        'Skaleeritav tehnoloogia infrastruktuur',
        'Tugev bränd',
        'Kasutajate lojaalsus'
    ])
    
    doc.add_heading('13.2 Investeeringu nõuded', 2)
    doc.add_paragraph('Esimese aasta investeering:')
    add_bullet_list(doc, [
        'Arendus: 7,000 € (juba tehtud)',
        'Turundus: 15,000 €',
        'Infrastruktuur: 3,300 €',
        'Personal: 1,500 €',
        'Muud kulud: 1,400 €',
        'Kokku: 28,200 €'
    ])
    
    doc.add_paragraph('Täiendavad investeeringud:')
    add_bullet_list(doc, [
        '2. aasta: 73,600 €',
        '3. aasta: 150,600 €',
        'Kokku 3 aastaks: 252,400 €'
    ])
    
    doc.add_heading('13.3 Oodatav tulemus', 2)
    doc.add_paragraph('Finantsiline tulemus:')
    add_bullet_list(doc, [
        '1. aasta: -10,725 € (investeering)',
        '2. aasta: +62,800 € (ROI: 223%)',
        '3. aasta: +226,400 € (ROI: 150%)',
        '3 aasta kogukasum: 278,475 €'
    ])
    
    doc.add_paragraph('Kasutajate kasv:')
    add_bullet_list(doc, [
        '1. aasta: 10,000–50,000 MAU',
        '2. aasta: 200,000–500,000 MAU',
        '3. aasta: 500,000–1,000,000 MAU'
    ])
    
    doc.add_paragraph('Tasuvuspunkt:')
    doc.add_paragraph('Kuu 14–16 (1. aasta lõpus)')
    
    doc.add_heading('13.4 Järgmised sammud', 2)
    doc.add_paragraph('Lähiajaline (1–3 kuud):')
    add_bullet_list(doc, [
        'Premium funktsioonide arendamine',
        'Google AdSense integratsioon',
        'Blogi käivitamine',
        'Sotsiaalmeedia kontode loomine',
        'Kasutajate tagasiside kogumine'
    ])
    
    doc.add_paragraph('Keskmine tähtaeg (4–12 kuud):')
    add_bullet_list(doc, [
        'Premium tellimuste käivitamine',
        'Partnerluste arendamine',
        'API arendamine',
        'B2B lahenduste arendamine',
        'Rahvusvaheline laienemine'
    ])
    
    doc.add_paragraph('Pikaajaline (2–3 aastat):')
    add_bullet_list(doc, [
        'Mobiilirakenduse arendamine',
        'Valge märgistus',
        'Ettevõtete lahendused',
        'Täiendavad tulude allikad',
        'Müügi võimalused'
    ])
    
    # Save document
    filename = 'FlipUnit_Ariplaan.docx'
    doc.save(filename)
    print(f"✅ Word dokument on edukalt loodud: {filename}")
    print(f"📄 Faili asukoht: {filename}")
    return filename


def create_tootukassa_business_plan():
    """Create business plan formatted for Estonian Unemployment Insurance Fund (Töötukassa)"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('ÄRIPLAAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_heading('FlipUnit.eu - Veebipõhine konverterite platvorm', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('Ettevõtlusega alustamise toetuse taotlus', style='Normal')
    doc.add_paragraph('Eesti Töötukassa', style='Normal')
    
    # Add page break
    doc.add_page_break()
    
    # ========== 1. ETTEVÕTTE KIRJELDUS ==========
    doc.add_heading('1. Ettevõtte kirjeldus', 1)
    
    doc.add_heading('1.1 Ettevõtte nimi ja tegevusala', 2)
    doc.add_paragraph('Ettevõtte nimi: FlipUnit.eu')
    doc.add_paragraph('Tegevusala: Veebipõhised tööriistad ja konverterid')
    doc.add_paragraph('Ettevõtte vorm: FIE (Füüsilisest isikust ettevõtja)')
    
    doc.add_heading('1.2 Ettevõtte eesmärgid', 2)
    doc.add_paragraph(
        'FlipUnit.eu on tasuta veebipõhine konverterite platvorm, mis pakub 100+ tööriista 12 kategoorias '
        'ilma registreerimise või sisselogimiseta. Platvorm on mõeldud kasutajatele, kes vajavad kiireid '
        'konversioone ilma tarkvara paigaldamise või kontode loomise vajaduseta.'
    )
    doc.add_paragraph('Peamised eesmärgid:')
    add_bullet_list(doc, [
        'Pakkuda kõike ühes kohas - ühe koha lahendus kõigile konversioonivajadustele',
        'Tagada kiire ja turvaline töötlemine ilma registreerimise vajaduseta',
        'Luua skaleeritav ärimudel reklaamitulu ja premium tellimuste kaudu',
        'Kasvada esimesel aastal 50,000–100,000 kuukülastajani',
        'Saavutada tasuvuspunkt 14–16 kuu jooksul'
    ])
    
    doc.add_heading('1.3 Toote/teenuse kirjeldus', 2)
    doc.add_paragraph(
        'FlipUnit.eu pakub laia valikut veebipõhiseid tööriistu, mis võimaldavad kasutajatel teha erinevaid '
        'konversioone ilma tarkvara paigaldamise vajaduseta. Platvorm on jagatud 12 kategooriasse:'
    )
    
    doc.add_heading('1.3.1 Mõõtühikute konverterid', 3)
    doc.add_paragraph('6 tööriista: pikkus, kaal, temperatuur, maht, pindala, kiirus')
    
    doc.add_heading('1.3.2 Pildi konverterid', 3)
    doc.add_paragraph('9 tööriista: formaadi konversioonid (JPEG, PNG, WebP, SVG), pildiredigeerimine (suuruse muutmine, pööramine, vesimärgid, jne)')
    
    doc.add_heading('1.3.3 Meedia konverterid', 3)
    doc.add_paragraph('9 tööriista: video/audio konversioonid, MP4 → MP3, video → GIF, audio tööriistad')
    
    doc.add_heading('1.3.4 PDF tööriistad', 3)
    doc.add_paragraph('11 tööriista: PDF ühendamine, jagamine, konversioonid, OCR, kompressioon')
    
    doc.add_heading('1.3.5 Muud kategooriad', 3)
    doc.add_paragraph('Valuuta ja krüpto konverter, arhiivi konverterid, teksti konverterid, arendaja tööriistad, utiliidid, värvi valija, veebisaidi staatuse kontrollija, YouTube pisipildi allalaadija')
    
    doc.add_paragraph('Kokku: 100+ tööriista 12 kategoorias')
    
    doc.add_heading('1.4 Põhiväärtus', 2)
    doc.add_paragraph('FlipUnit.eu eristub konkurentidest järgmiste omaduste poolest:')
    add_bullet_list(doc, [
        'Laiem tööriistade valik (100+ vs 20–50 konkurentidel)',
        'Registreerimise vajadus puudub - väiksem takistus kasutajatele',
        'Kiirem töötlemine tänu optimeeritud algoritmidele',
        'Privaatsus - failid kustutatakse automaatselt',
        'Mobiilne optimeerimine - töötab kõigil seadmetel',
        'SEO optimeerimine - tugev otsingunähtavus'
    ])
    
    # ========== 2. ETTEVÕTJA ANDMED ==========
    doc.add_heading('2. Ettevõtja andmed', 1)
    
    doc.add_heading('2.1 Isikuandmed', 2)
    doc.add_paragraph('(Isikuandmed lisatakse eraldi CV-sse, mis lisatakse taotlusele)')
    doc.add_paragraph('Põhjalikuma info oma CV-sse, mille lisan taotlusele.')
    
    doc.add_heading('2.2 Haridus ja kogemus', 2)
    doc.add_paragraph('(Detailne kirjeldus CV-s)')
    doc.add_paragraph('Panen põhjalikuma info oma CV-sse, mille lisan taotlusele.')
    
    doc.add_heading('2.3 Tugevused ja nõrkused', 2)
    doc.add_paragraph('Äriplaani koostamise käigus saan läbi mõelda enda tugevused ja nõrkused ning äri põhialused.')
    
    doc.add_heading('2.3.1 Tugevused', 3)
    add_bullet_list(doc, [
        'Tehniline oskus: Django, Python, veebiarendus',
        'Kogemus projektide arendamises',
        'Süsteemne mõtlemine ja probleemide lahendamine',
        'Motivatsioon ja pühendumus projekti elluviimisele'
    ])
    
    doc.add_heading('2.3.2 Arendamist vajavad oskused', 3)
    add_bullet_list(doc, [
        'Turundus- ja müügioskused',
        'Äriplaneerimise kogemus',
        'Kliendisuhtluse kogemus',
        'Finantsjuhtimise oskused'
    ])
    
    doc.add_heading('2.4 Mõtteviis', 2)
    doc.add_paragraph('Äriplaani koostamise võttel mõtlen suurelt ja samas hindan realistlikult, mida ma suudan ellu viia.')
    doc.add_paragraph('Kui äriplaan valmis saan, siis tean, kuhu võin oma äriga tulevikus välja jõuda (nt mida, kellele, kuidas).')
    
    # ========== 3. TURU ANALÜÜS ==========
    doc.add_heading('3. Turu analüüs', 1)
    
    doc.add_heading('3.1 Sihtrühmad', 2)
    
    doc.add_heading('3.1.1 Peamised kasutajad', 3)
    doc.add_paragraph('Üldkasutajad:')
    add_bullet_list(doc, [
        'Pildi- ja PDF-konversioonid',
        'Dokumentide töötlemine',
        'Igapäevased konversioonid'
    ])
    doc.add_paragraph('Õpilased ja õpetajad:')
    add_bullet_list(doc, [
        'Mõõtühikute konversioonid',
        'Dokumentide tööriistad',
        'Õppematerjalide töötlemine'
    ])
    doc.add_paragraph('Väikeettevõtjad:')
    add_bullet_list(doc, [
        'Dokumentide töötlemine',
        'Pildiredigeerimine',
        'Meediafailide konversioonid'
    ])
    
    doc.add_heading('3.1.2 Teisene sihtrühm', 3)
    doc.add_paragraph('Veebiarendajad, sisuloojad, professionaalid')
    
    doc.add_heading('3.2 Turu suurus', 2)
    doc.add_paragraph('Globaalne turu ülevaade:')
    add_bullet_list(doc, [
        'Veebipõhiste tööriistade turg: hinnanguliselt 5–10 miljardit eurot (kasv 15–20% aastas)',
        'Konverterite turg: kõrge otsingumaht (miljonid päevas)',
        'Registreerimiseta tööriistad: kasvav nõudlus privaatsuse tõttu'
    ])
    
    doc.add_paragraph('Eesti turu potentsiaal:')
    add_bullet_list(doc, [
        'Eesti elanikkond: ~1,3 miljonit',
        'Internetikasutajad: ~95% (1,2 miljonit)',
        'Potentsiaalne kasutajabaas: 50,000–100,000 aktiivset kasutajat'
    ])
    
    doc.add_paragraph('Rahvusvaheline turu potentsiaal:')
    add_bullet_list(doc, [
        'Euroopa: 500+ miljonit internetikasutajat',
        'Põhjusmärksõnad: miljonid päevas',
        'Potentsiaalne kasutajabaas: 1–5 miljonit kuus'
    ])
    
    doc.add_heading('3.3 Konkurentsianalüüs', 2)
    
    doc.add_heading('3.3.1 Peamised konkurendid', 3)
    doc.add_paragraph('1. Zamzar - lai valik, API, aga registreerimine ja aeglane töötlemine')
    doc.add_paragraph('2. CloudConvert - API, pilve integratsioonid, aga keerukas kasutajaliides ja kõrged hinnad')
    doc.add_paragraph('3. Convertio - lihtne kasutajaliides, aga piiratud valik ja reklaamid')
    doc.add_paragraph('4. Online-Convert - lai valik, aga registreerimine ja aeglane töötlemine')
    
    doc.add_heading('3.3.2 Konkurentsieelised', 3)
    add_bullet_list(doc, [
        'Laiem tööriistade valik (100+ vs 20–50)',
        'Registreerimise vajadus puudub',
        'Kiirem töötlemine',
        'Parem kasutajakogemus',
        'Privaatsus (failid kustutatakse automaatselt)',
        'Mobiilne optimeerimine',
        'SEO optimeerimine'
    ])
    
    # ========== 4. TURUNDUSSTRATEEGIA ==========
    doc.add_heading('4. Turundusstrateegia', 1)
    
    doc.add_heading('4.1 Turunduse eesmärgid', 2)
    doc.add_paragraph('Esimese aasta eesmärgid:')
    add_bullet_list(doc, [
        '10,000–25,000 külastajat/kuu (esimesed 6 kuud)',
        '50,000–100,000 külastajat/kuu (aasta lõpus)',
        '1,000–5,000 aktiivset kasutajat/kuu',
        'Tugev SEO positsioon valitud märksõnadele'
    ])
    
    doc.add_heading('4.2 Turunduskanalid', 2)
    
    doc.add_heading('4.2.1 SEO ja sisuturundus', 3)
    doc.add_paragraph('Peamine fookus:')
    add_bullet_list(doc, [
        'XML sitemap ja meta sildid',
        'Blogi: tööriistade juhendid',
        'Põhjusmärksõnade sihtimine',
        'Sisemine linkimine',
        'Kohalik SEO (Eesti)',
        'Mitmekeelne sisu (eesti, inglise, vene)'
    ])
    
    doc.add_heading('4.2.2 Sotsiaalmeedia turundus', 3)
    add_bullet_list(doc, [
        'YouTube: tööriistade demo videod ja "Kuidas" juhendid',
        'TikTok/Instagram Reels: lühikesed demo videod',
        'Twitter/X: uuendused ja näpunäited',
        'Reddit: kogukondade osalemine',
        'Product Hunt: platvormi käivitamine'
    ])
    
    doc.add_heading('4.2.3 Partnerlused', 3)
    add_bullet_list(doc, [
        'Integratsioonid tootlikkustööriistadega',
        'Arendajate kogukonnad',
        'Ülikoolid ja koolid',
        'Väikeettevõtted'
    ])
    
    doc.add_heading('4.2.4 Tasustatud reklaam', 3)
    doc.add_paragraph('Planeeritud kulud: 1,500–3,000 €/kuu esimesel aastal')
    add_bullet_list(doc, [
        'Google Ads: kõrge kavatsusega märksõnad',
        'Facebook/Instagram reklaam: sihtrühmade reklaamid',
        'Retargeting: tagasipöördumise kampaaniad'
    ])
    
    # ========== 5. ÄRIMUDEL JA TULUDE ALLIKAD ==========
    doc.add_heading('5. Ärimudel ja tulude allikad', 1)
    
    doc.add_heading('5.1 Praegune mudel', 2)
    doc.add_paragraph('Freemium mudel: tasuta juurdepääs kõigile tööriistadele ilma registreerimiseta')
    
    doc.add_heading('5.2 Tulude allikad', 2)
    
    doc.add_heading('5.2.1 Reklaamitulu (peamine)', 3)
    doc.add_paragraph('Google AdSense ja partnerreklaamid')
    doc.add_paragraph('Eeldatav tulu esimesel aastal: 12,000–36,000 €')
    
    doc.add_heading('5.2.2 Premium tellimused', 3)
    doc.add_paragraph('Premium funktsioonid: suuremad failid, prioriteetne töötlemine, API juurdepääs, ilma reklaamideta')
    doc.add_paragraph('Hind: 4,99–9,99 €/kuu')
    doc.add_paragraph('Eeldatav tulu esimesel aastal: 3,600–7,200 €')
    
    doc.add_heading('5.2.3 B2B lahendused', 3)
    doc.add_paragraph('API juurdepääs ettevõtetele, kohandatud integratsioonid')
    doc.add_paragraph('Eeldatav tulu esimesel aastal: 1,200–3,000 €')
    
    doc.add_heading('5.3 Eeldatav kogutulu', 2)
    doc.add_paragraph('Esimese aasta kogutulu: 18,000–52,200 €')
    
    # ========== 6. FINANTSPROGNOOSID ==========
    doc.add_heading('6. Finantsprognoosid', 1)
    
    doc.add_heading('6.1 Esimese aasta finantsprognoos', 2)
    
    headers_q1 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q1 = [
        ["Reklaamitulu", "600 €", "1,800 €", "3,600 €", "6,000 €", "12,000 €"],
        ["Premium tellimused", "0 €", "525 €", "1,050 €", "2,100 €", "3,675 €"],
        ["Partnerreklaamid", "300 €", "400 €", "500 €", "600 €", "1,800 €"],
        ["Kogutulu", "900 €", "2,725 €", "5,150 €", "8,700 €", "17,475 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "600 €", "750 €", "900 €", "1,050 €", "3,300 €"],
        ["Turundus", "1,500 €", "3,000 €", "4,500 €", "6,000 €", "15,000 €"],
        ["Arendus", "1,000 €", "1,500 €", "2,000 €", "2,500 €", "7,000 €"],
        ["Personal (osaline)", "0 €", "0 €", "500 €", "1,000 €", "1,500 €"],
        ["Muud kulud", "200 €", "300 €", "400 €", "500 €", "1,400 €"],
        ["Kogukulud", "3,300 €", "5,550 €", "8,300 €", "11,050 €", "28,200 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "-2,400 €", "-2,825 €", "-3,150 €", "-2,350 €", "-10,725 €"]
    ]
    
    add_table_with_data(doc, headers_q1, rows_q1)
    doc.add_paragraph('Tasuvuspunkt: kuu 14–16')
    
    doc.add_heading('6.2 Teise aasta finantsprognoos', 2)
    
    headers_q2 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q2 = [
        ["Reklaamitulu", "15,000 €", "18,000 €", "21,000 €", "24,000 €", "78,000 €"],
        ["Premium tellimused", "6,000 €", "7,500 €", "9,000 €", "10,500 €", "33,000 €"],
        ["B2B lahendused", "3,000 €", "4,500 €", "6,000 €", "7,500 €", "21,000 €"],
        ["Partnerreklaamid", "800 €", "1,000 €", "1,200 €", "1,400 €", "4,400 €"],
        ["Kogutulu", "24,800 €", "31,000 €", "37,200 €", "43,400 €", "136,400 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "1,200 €", "1,400 €", "1,600 €", "1,800 €", "6,000 €"],
        ["Turundus", "7,500 €", "9,000 €", "10,500 €", "12,000 €", "39,000 €"],
        ["Arendus", "3,000 €", "3,500 €", "4,000 €", "4,500 €", "15,000 €"],
        ["Personal", "2,000 €", "2,500 €", "3,000 €", "3,500 €", "11,000 €"],
        ["Muud kulud", "500 €", "600 €", "700 €", "800 €", "2,600 €"],
        ["Kogukulud", "14,200 €", "17,000 €", "19,800 €", "22,600 €", "73,600 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "10,600 €", "14,000 €", "17,400 €", "20,800 €", "62,800 €"]
    ]
    
    add_table_with_data(doc, headers_q2, rows_q2)
    doc.add_paragraph('Kasumimarginaal: 46%')
    
    doc.add_heading('6.3 Kolmanda aasta finantsprognoos', 2)
    
    headers_q3 = ["Näitaja", "Q1", "Q2", "Q3", "Q4", "Aasta kokku"]
    rows_q3 = [
        ["Reklaamitulu", "45,000 €", "50,000 €", "55,000 €", "60,000 €", "210,000 €"],
        ["Premium tellimused", "18,000 €", "21,000 €", "24,000 €", "27,000 €", "90,000 €"],
        ["B2B lahendused", "12,000 €", "15,000 €", "18,000 €", "21,000 €", "66,000 €"],
        ["Partnerreklaamid", "2,000 €", "2,500 €", "3,000 €", "3,500 €", "11,000 €"],
        ["Kogutulu", "77,000 €", "88,500 €", "100,000 €", "111,500 €", "377,000 €"],
        ["", "", "", "", "", ""],
        ["Infrastruktuur", "2,400 €", "2,700 €", "3,000 €", "3,300 €", "11,400 €"],
        ["Turundus", "15,000 €", "18,000 €", "21,000 €", "24,000 €", "78,000 €"],
        ["Arendus", "6,000 €", "7,000 €", "8,000 €", "9,000 €", "30,000 €"],
        ["Personal", "5,000 €", "6,000 €", "7,000 €", "8,000 €", "26,000 €"],
        ["Muud kulud", "1,000 €", "1,200 €", "1,400 €", "1,600 €", "5,200 €"],
        ["Kogukulud", "29,400 €", "34,900 €", "40,400 €", "45,900 €", "150,600 €"],
        ["", "", "", "", "", ""],
        ["Puhaskasum", "47,600 €", "53,600 €", "59,600 €", "65,600 €", "226,400 €"]
    ]
    
    add_table_with_data(doc, headers_q3, rows_q3)
    doc.add_paragraph('Kasumimarginaal: 60%')
    
    # ========== 7. INVESTEERINGU VAJADUS ==========
    doc.add_heading('7. Investeeringu vajadus', 1)
    
    doc.add_heading('7.1 Esimese aasta investeeringu nõuded', 2)
    doc.add_paragraph('Esimese aasta investeering:')
    add_bullet_list(doc, [
        'Arendus: 7,000 € (juba tehtud)',
        'Turundus: 15,000 €',
        'Infrastruktuur: 3,300 €',
        'Personal: 1,500 €',
        'Muud kulud: 1,400 €',
        'Kokku: 28,200 €'
    ])
    
    doc.add_heading('7.2 Töötukassa toetuse kasutamine', 2)
    doc.add_paragraph('Töötukassa toetust kasutatakse järgmiste eesmärkide jaoks:')
    add_bullet_list(doc, [
        'Turunduskulud: SEO optimeerimine, sisuturundus, sotsiaalmeedia reklaam',
        'Infrastruktuuri kulud: server hosting, CDN, andmebaas',
        'Arenduskulud: täiendavad funktsioonid, turundusmaterjalide arendamine',
        'Personal: osalise tööajaga turunduse spetsialisti palkamine'
    ])
    
    doc.add_heading('7.3 Oodatav tulemus', 2)
    doc.add_paragraph('Investeeringu abil saavutatakse:')
    add_bullet_list(doc, [
        'Platvormi stabiliseerimine ja kasutajate kasv',
        'Tugev SEO positsioon',
        'Premium funktsioonide käivitamine',
        'Reklaamitulu käivitamine',
        'Tasuvuspunkt 14–16 kuu jooksul'
    ])
    
    # ========== 8. RISKIANALÜÜS ==========
    doc.add_heading('8. Riskianalüüs', 1)
    
    doc.add_heading('8.1 Peamised riskid', 2)
    
    doc.add_heading('8.1.1 Kõrged infrastruktuuri kulud', 3)
    doc.add_paragraph('Risk: Kasutajate kasv suurendab serverikulusid')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine: Tõhus töötlemise algoritm, pilve auto-skaleerimine, CDN kasutamine')
    
    doc.add_heading('8.1.2 Tugev konkurents', 3)
    doc.add_paragraph('Risk: Suuremad konkurendid võivad konkureerida')
    doc.add_paragraph('Tõenäosus: Kõrge')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine: Ainulaadsed funktsioonid, parem kasutajakogemus, kiire arendus')
    
    doc.add_heading('8.1.3 Juridilised/järgimisnõuded', 3)
    doc.add_paragraph('Risk: GDPR, andmekaitse, autoriõigused')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Kõrge')
    doc.add_paragraph('Leevendamine: Privaatsuskeskne disain, selged kasutustingimused, juridiline nõustamine')
    
    doc.add_heading('8.1.4 Turunduse ebaõnnestumine', 3)
    doc.add_paragraph('Risk: Liikluse kasv on aeglasem kui oodatud')
    doc.add_paragraph('Tõenäosus: Keskmine')
    doc.add_paragraph('Mõju: Keskmine')
    doc.add_paragraph('Leevendamine: Mitmekesine turundusstrateegia, SEO fookus, sotsiaalmeedia kohalolek')
    
    # ========== 9. TÖÖKOHTADE LOOMINE ==========
    doc.add_heading('9. Töökohtade loomine', 1)
    
    doc.add_heading('9.1 Lähiajaline (1. aasta)', 2)
    doc.add_paragraph('Esimesel aastal:')
    add_bullet_list(doc, [
        'Ettevõtja ise: täisajaline töö',
        'Osalise tööajaga turunduse spetsialist: 0,5 kohta (aasta teisel poolel)'
    ])
    
    doc.add_heading('9.2 Keskmine tähtaeg (2. aasta)', 2)
    doc.add_paragraph('Teisel aastal:')
    add_bullet_list(doc, [
        'Ettevõtja ise: täisajaline töö',
        'Täisajaline arendaja: 1 koht',
        'Täisajaline turunduse spetsialist: 1 koht',
        'B2B müügispetsialist: 1 koht',
        'Klienditugi spetsialist: 1 koht',
        'Kokku: 5 töökohta'
    ])
    
    doc.add_heading('9.3 Pikaajaline (3. aasta)', 2)
    doc.add_paragraph('Kolmandal aastal:')
    add_bullet_list(doc, [
        'Arendusmeeskond: 3–5 inimest',
        'Turundusmeeskond: 2–3 inimest',
        'Müügimeeskond: 2–3 inimest',
        'Toetusmeeskond: 2–3 inimest',
        'Juhtimine: CEO, CTO, CMO',
        'Kokku: 10–15 töökohta'
    ])
    
    # ========== 10. KOKKUVÕTE ==========
    doc.add_heading('10. Kokkuvõte', 1)
    
    doc.add_paragraph(
        'FlipUnit.eu on positsioneeritud kasutama veebipõhiste tööriistade turu kasvu. '
        '100+ tööriistaga, registreerimiseta mudeliga ja skaleeritava arhitektuuriga on platvorm '
        'valmis kasvama tasuta utiliidist tasuliseks platvormiks reklaamitulu, premium tellimuste '
        'ja ettevõtete lahenduste kaudu.'
    )
    
    doc.add_heading('10.1 Peamised edu tegurid', 2)
    add_bullet_list(doc, [
        'Tugev SEO alus',
        'Kasutajakeskne disain',
        'Pidev funktsioonide arendus',
        'Tõhus tasustamisstrateegia',
        'Skaleeritav tehnoloogia infrastruktuur',
        'Tugev bränd',
        'Kasutajate lojaalsus'
    ])
    
    doc.add_heading('10.2 Oodatav tulemus', 2)
    doc.add_paragraph('Finantsiline tulemus:')
    add_bullet_list(doc, [
        '1. aasta: -10,725 € (investeering)',
        '2. aasta: +62,800 € (ROI: 223%)',
        '3. aasta: +226,400 € (ROI: 150%)',
        '3 aasta kogukasum: 278,475 €'
    ])
    
    doc.add_paragraph('Kasutajate kasv:')
    add_bullet_list(doc, [
        '1. aasta: 10,000–50,000 MAU',
        '2. aasta: 200,000–500,000 MAU',
        '3. aasta: 500,000–1,000,000 MAU'
    ])
    
    doc.add_paragraph('Tasuvuspunkt: Kuu 14–16 (1. aasta lõpus)')
    
    doc.add_heading('10.3 Järgmised sammud', 2)
    doc.add_paragraph('Lähiajaline (1–3 kuud):')
    add_bullet_list(doc, [
        'Premium funktsioonide arendamine',
        'Google AdSense integratsioon',
        'Blogi käivitamine',
        'Sotsiaalmeedia kontode loomine',
        'Kasutajate tagasiside kogumine'
    ])
    
    # Save document
    filename = 'FlipUnit_Ariplaan_Tootukassa.docx'
    doc.save(filename)
    print(f"✅ Töötukassa äriplaan on edukalt loodud: {filename}")
    print(f"📄 Faili asukoht: {filename}")
    return filename


if __name__ == "__main__":
    try:
        print("Valige äriplaani tüüp:")
        print("1. Tavaline äriplaan")
        print("2. Töötukassa äriplaan")
        choice = input("Sisestage valik (1 või 2): ").strip()
        
        if choice == "2":
            create_tootukassa_business_plan()
        else:
            create_business_plan_document()
    except Exception as e:
        print(f"❌ Viga dokumendi loomisel: {e}")
        import traceback
        traceback.print_exc()

